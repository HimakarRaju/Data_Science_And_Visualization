import logging
import os
import pickle
import re
import shutil

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
from docx import Document
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    RandomForestRegressor,
    GradientBoostingRegressor,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.metrics import (
    confusion_matrix,
    classification_report,
    mean_absolute_error,
    mean_squared_error,
    r2_score,
)
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.svm import SVC, SVR
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor

# Setup Plotly renderer
pio.renderers.default = "browser"

# Global analysis log
analysis_log = {
    "steps": [],
    "reasons": [],
}


def log_step(step, reason):
    analysis_log["steps"].append(step)
    analysis_log["reasons"].append(reason)


# Setup logging
def setup_logging(log_file):
    logging.basicConfig(
        filename=log_file,
        level=logging.ERROR,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


# Check for symbols and spaces in column names
def check_column_names(df):
    """Check for symbols and spaces in DataFrame column names."""
    invalid_columns = [col for col in df.columns if re.search(r"[^a-zA-Z0-9_]", col)]
    if invalid_columns:
        print(f"Warning: Columns with symbols/spaces detected: {invalid_columns}")
        df.columns = df.columns.str.replace(r"[^a-zA-Z0-9_]", "_", regex=True)
    return df


# Create required output directories
def create_directories(output_directory, dataset_name):
    dirs = ["analysis", "plots", "time_series"]
    created_dirs = {}
    for d in dirs:
        created_dirs[d] = os.path.join(output_directory, dataset_name, d)
        os.makedirs(created_dirs[d], exist_ok=True)
    return created_dirs


# Load target column data from pickle
def load_target_from_pickle(pkl_dir):
    pkl_file = os.path.join(pkl_dir, "target_column.pkl")
    if os.path.exists(pkl_file):
        with open(pkl_file, "rb") as f:
            target_column = pickle.load(f)
        print(f"Previous target column found: {target_column}")
        return target_column
    return None


# Save target column to pickle
def save_target_to_pickle(target_column, pkl_dir):
    pkl_file = os.path.join(pkl_dir, "target_column.pkl")
    with open(pkl_file, "wb") as f:
        pickle.dump(target_column, f)


# Save all data to pickle
def save_all_data_to_pickle(df, pkl_dir):
    pkl_file = os.path.join(pkl_dir, "data.pkl")
    with open(pkl_file, "ab") as f:
        pickle.dump(df, f)


# Suggest target column
def suggest_target_column(df, pkl_dir):
    log_step(
        "Suggesting target column based on existing columns",
        "To identify the appropriate target for modeling.",
    )
    target_column = load_target_from_pickle(pkl_dir)
    if not target_column:
        for column in df.columns:
            if any(x in column.lower() for x in ["price", "target", "label"]):
                target_column = column
                break
        if not target_column:
            target_column = input(
                f"No suitable target column found. Choose from: {list(df.columns)} "
            )

    print(f"Suggested target column: {target_column}")
    confirmation = input("Is this okay? (y/n): ")
    if confirmation.lower() != "y":
        target_column = input(
            f"Please choose the target column from: {list(df.columns)} "
        )

    save_target_to_pickle(target_column, pkl_dir)
    return target_column


# Data transformation
def transform_data(df):
    log_step(
        "Transforming data by imputing and normalizing",
        "To handle missing values and ensure data consistency.",
    )
    print("Starting data transformation...")
    imputer = SimpleImputer(strategy="mean")
    df_imputed = df.copy()

    numerical_columns = df.select_dtypes(include=[np.number]).columns
    df_imputed[numerical_columns] = imputer.fit_transform(df_imputed[numerical_columns])

    # Normalization
    if df_imputed[numerical_columns].max().max() != 0:
        normalizer = Normalizer()
        df_imputed[numerical_columns] = normalizer.fit_transform(
            df_imputed[numerical_columns]
        )

    # Label encoding for categorical variables
    label_encoders = {}
    categorical_columns = df.select_dtypes(include=["object"]).columns
    for column in categorical_columns:
        le = LabelEncoder()
        df_imputed[column] = le.fit_transform(df_imputed[column].astype(str))
        label_encoders[column] = le

    print("Data transformation completed.")
    return df_imputed, label_encoders


# Evaluate models
def evaluate_models(df, target_column, plots_dir):
    log_step(
        "Evaluating models based on the target column type",
        "To assess performance using suitable algorithms.",
    )
    results = {}
    X = df.drop(columns=[target_column])
    y = df[target_column]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )

    models = select_models(y)

    for model_name, model in models.items():
        try:
            model.fit(X_train, y_train)
            predictions = model.predict(X_test)
            results[model_name] = evaluate_model(
                y, y_test, predictions, model_name, plots_dir
            )
        except Exception as e:
            log_error(model_name, e)

    return results


def select_models(y):
    if y.dtype in [np.int64, np.float64] and len(y.unique()) > 2:
        return {  # Regression
            "Linear Regression": LinearRegression(),
            "Random Forest Regressor": RandomForestRegressor(),
            "Gradient Boosting Regressor": GradientBoostingRegressor(),
            "K-Neighbors Regressor": KNeighborsRegressor(),
            "Support Vector Regressor": SVR(),
            "Decision Tree Regressor": DecisionTreeRegressor(),
        }
    else:  # Classification
        return {
            "Logistic Regression": LogisticRegression(max_iter=1000),
            "Random Forest": RandomForestClassifier(),
            "Gradient Boosting": GradientBoostingClassifier(),
            "K-Neighbors": KNeighborsClassifier(),
            "Support Vector Classifier": SVC(),
            "Naive Bayes": GaussianNB(),
            "Decision Tree": DecisionTreeClassifier(),
        }


def evaluate_model(y, y_test, predictions, model_name, plots_dir):
    if y.dtype in [np.int64, np.float64] and len(y.unique()) > 2:  # Regression
        return regression_metrics(y_test, predictions, model_name)
    else:  # Classification
        return classification_metrics(y_test, predictions, model_name, plots_dir)


def regression_metrics(y_test, predictions, model_name):
    mae = mean_absolute_error(y_test, predictions)
    mse = mean_squared_error(y_test, predictions)
    r2 = r2_score(y_test, predictions)

    print(f"{model_name} Metrics:\nMAE: {mae}, MSE: {mse}, R²: {r2}")
    return {"MAE": mae, "MSE": mse, "R²": r2}


def classification_metrics(y_test, predictions, model_name, plots_dir):
    conf_matrix = confusion_matrix(y_test, predictions)
    class_report = classification_report(y_test, predictions)

    print(f"Confusion Matrix for {model_name}:\n{conf_matrix}")
    print(f"Classification Report for {model_name}:\n{class_report}")

    # Plotting the confusion matrix with Plotly
    fig = px.imshow(
        conf_matrix,
        text_auto=True,
        title=f"Confusion Matrix: {model_name}",
        labels=dict(x="Predicted Label", y="True Label"),
        color_continuous_scale="Blues",
    )
    fig.write_html(os.path.join(plots_dir, f"confusion_matrix_{model_name}.html"))

    return {"Confusion Matrix": conf_matrix, "Classification Report": class_report}


# Generate and save plots
def generate_and_save_plots(df, plots_dir, target_column):
    log_step(
        "Generating and saving plots for data analysis",
        "To visualize data distributions and relationships.",
    )
    print("Generating plots...")
    generate_histograms(df, plots_dir)
    generate_correlation_matrix(df, plots_dir)
    generate_target_feature_plots(df, target_column, plots_dir)


def generate_histograms(df, plots_dir):
    numerical_cols = df.select_dtypes(include=[np.number]).columns
    for col in numerical_cols:
        fig = px.histogram(
            df,
            x=col,
            title=f"Histogram of {col}",
            labels={col: col},
        )
        fig.write_html(os.path.join(plots_dir, f"histogram_{col}.html"))


def generate_correlation_matrix(df, plots_dir):
    corr_matrix = df.corr()
    fig = px.imshow(
        corr_matrix, title="Correlation Matrix", color_continuous_scale="Viridis"
    )
    fig.write_html(os.path.join(plots_dir, "correlation_matrix.html"))


def generate_target_feature_plots(df, target_column, plots_dir):
    for col in df.drop(columns=[target_column]).columns:
        fig = px.scatter(
            df,
            x=col,
            y=target_column,
            title=f"{target_column} vs {col}",
            labels={col: col, target_column: target_column},
        )
        fig.write_html(os.path.join(plots_dir, f"{target_column}_vs_{col}.html"))


# Move generated files to their respective folders
def move_files_to_respective_folders(df, plots_dir, analysis_dir, pkl_dir):
    # Move HTML files to plots folder
    for filename in os.listdir(plots_dir):
        if filename.endswith(".html"):
            shutil.move(
                os.path.join(plots_dir, filename), os.path.join(plots_dir, filename)
            )

    # Move DOCX file to analysis folder
    docx_file_path = os.path.join(analysis_dir, "Analysis_Report.docx")
    if os.path.exists(docx_file_path):
        shutil.move(docx_file_path, os.path.join(analysis_dir, "Analysis_Report.docx"))

    # Move data and target column to the specified directory
    save_all_data_to_pickle(df, pkl_dir)  # Now df is accessible here


def log_error(model_name, error):
    logging.error(f"Error in model {model_name}: {error}")


# Determine the best model based on metrics
def determine_best_model(results):
    best_model = None
    best_score = float("-inf")
    for model_name, metrics in results.items():
        if "R²" in metrics and metrics["R²"] > best_score:
            best_score = metrics["R²"]
            best_model = model_name
    return {"name": best_model, "R²": best_score}


# Finalize analysis and create report
def finalize_analysis(results, plots_dir, analysis_dir):
    print("Analysis Steps and Reasons:")
    for step, reason in zip(analysis_log["steps"], analysis_log["reasons"]):
        print(f"Step: {step}\nReason: {reason}\n")

    # Create Word document
    doc = Document()
    doc.add_heading("Analysis Report", level=1)

    # Methods section
    doc.add_heading("Methods Used", level=2)
    for step, reason in zip(analysis_log["steps"], analysis_log["reasons"]):
        doc.add_paragraph(f"{step}: {reason}")

    # Results section
    doc.add_heading("Model Results", level=2)
    for model_name, metrics in results.items():
        doc.add_heading(model_name, level=3)
        if "MAE" in metrics:
            doc.add_paragraph(f"MAE: {metrics['MAE']:.4f}")
            doc.add_paragraph(f"MSE: {metrics['MSE']:.4f}")
            doc.add_paragraph(f"R²: {metrics['R²']:.4f}")
        else:
            doc.add_paragraph(metrics["Classification Report"])
        doc.add_paragraph(
            f"Confusion Matrix: See {os.path.join(plots_dir, f'confusion_matrix_{model_name}.html')}"
        )

    # Best model section
    best_model_info = determine_best_model(results)
    doc.add_heading("Best Model", level=2)
    doc.add_paragraph(
        f"The best model is: {best_model_info['name']} with R²: {best_model_info['R²']:.4f}"
    )

    # Save the document in the analysis folder
    doc_file_path = os.path.join(analysis_dir, "Analysis_Report.docx")
    doc.save(doc_file_path)

    print(f"Report saved to {doc_file_path}")


# Run the analysis
def run_analysis(file_path, output_directory, pikl_folder):
    print(f"Loading data from {file_path}...")
    df = pd.read_csv(file_path)

    df = check_column_names(df)
    dataset_name = os.path.splitext(os.path.basename(file_path))[0]
    dirs = create_directories(output_directory, dataset_name)
    pkl_dir = pikl_folder

    target_column = suggest_target_column(df, pkl_dir)
    df_transformed, label_encoders = transform_data(df)
    save_all_data_to_pickle(df_transformed, pkl_dir)

    results = evaluate_models(df_transformed, target_column, dirs["plots"])
    generate_and_save_plots(df_transformed, dirs["plots"], target_column)

    # Make sure to pass df when calling this function
    move_files_to_respective_folders(df, dirs["plots"], dirs["analysis"], pkl_dir)
    finalize_analysis(results, dirs["plots"], dirs["analysis"])


if __name__ == "__main__":
    run_analysis(
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\Datasets\\user_behavior_dataset.csv",
        # Replace with your data file path
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts1",
        # Replace with your desired output directory
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\pkls",
        # Replace with your desired pkl output directory
    )
