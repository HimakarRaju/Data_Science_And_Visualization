import logging
import os
import pickle
import re
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    RandomForestRegressor,
    GradientBoostingRegressor,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.metrics import (
    confusion_matrix,
    classification_report,
    mean_absolute_error,
    mean_squared_error,
    r2_score,
)
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.svm import SVC, SVR
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor

# Setup Plotly renderer
pio.renderers.default = "browser"

# Global analysis log
analysis_log = {"steps": [], "reasons": []}


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a Word document."""
    run = paragraph.add_run(text)
    r_id = paragraph.part.relate_to(url, "hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set("r:id", r_id)
    run._r.append(hyperlink)
    return paragraph


def check_column_names(df):
    """Check and standardize column names."""
    df.columns = [re.sub(r"\s+", "_", col).strip().lower() for col in df.columns]
    return df


def setup_logging(log_filename):
    """Setup logging configuration."""
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[logging.FileHandler(log_filename), logging.StreamHandler()],
    )


def log_step(step, reason):
    """Log the analysis steps and reasons."""
    analysis_log["steps"].append(step)
    analysis_log["reasons"].append(reason)


def save_analysis_log(
    log_data,
    output_dir,
    models_info,
    best_model,
    plots_dir,
    target_column,
    filename="analysis_log.docx",
):
    """Generate and save the analysis log document."""
    doc = Document()
    doc.add_heading("Data Analysis Log", 0)

    doc.add_heading("Analysis Steps and Reasons", level=1)
    for i, (step, reason) in enumerate(zip(log_data["steps"], log_data["reasons"])):
        doc.add_paragraph(f"{i + 1}. {step}: {reason}")

    doc.add_heading("Model Evaluation", level=1)
    for model_name, model_info in models_info.items():
        doc.add_heading(f"Model: {model_name}", level=2)
        doc.add_paragraph(f"Metrics:\n{model_info}")
        if model_info.get("Confusion Matrix"):
            confusion_matrix_plot = os.path.join(
                plots_dir, f"confusion_matrix_{model_name}.html"
            )
            doc.add_paragraph(f"Confusion Matrix for {model_name}:")
            add_hyperlink(
                doc.add_paragraph(),
                confusion_matrix_plot,
                f"Click here to view the Confusion Matrix plot for {model_name}",
            )

    doc.add_heading("Best Model", level=1)
    doc.add_paragraph(
        f"The best-performing model was {best_model} based on the evaluation metrics."
    )

    if best_model:
        correlation_matrix_plot = os.path.join(plots_dir, "correlation_matrix.html")
        add_hyperlink(
            doc.add_paragraph(),
            correlation_matrix_plot,
            "Click here to view the Correlation Matrix",
        )

        doc.add_heading(f"Scatter Plots for {target_column} vs Features", level=2)
        for col in models_info.get(best_model, {}).get("features", []):
            scatter_plot = os.path.join(
                plots_dir, f"scatter_{col}_vs_{target_column}.html"
            )
            add_hyperlink(
                doc.add_paragraph(),
                scatter_plot,
                f"Click here to view Scatter Plot for {col} vs {target_column}",
            )

    doc.save(os.path.join(output_dir, filename))


def create_directories(output_directory, dataset_name):
    """Create necessary directories for saving outputs."""
    dirs = ["analysis", "plots", "plotSVGs"]
    created_dirs = {d: os.path.join(output_directory, dataset_name, d) for d in dirs}
    for dir_path in created_dirs.values():
        os.makedirs(dir_path, exist_ok=True)
    return created_dirs


def load_pickle_data(pkl_dir):
    """Load data from multiple pickle files in the specified directory."""
    data = {}
    for filename in os.listdir(pkl_dir):
        if filename.endswith(".pkl"):
            with open(os.path.join(pkl_dir, filename), "rb") as f:
                try:
                    data.update(pickle.load(f))
                except EOFError:
                    continue
    return data


def save_to_pickle(
    df,
    pkl_dir,
    target_column,
    best_model=None,
    vectorizer=None,
    transformation_steps=None,
    dataset_name=None,
):
    """Save the dataset and transformation details to pickle."""
    dataset_metadata = {
        "target_column": target_column,
        "shape": df.shape,
        "columns": df.columns.tolist(),
        "best_model": best_model,
        "vectorizer": vectorizer,
        "transformation_steps": transformation_steps,
    }
    pkl_file = os.path.join(pkl_dir, f"{dataset_name}_data.pkl")
    with open(pkl_file, "wb") as f:
        pickle.dump(dataset_metadata, f)


def suggest_target_column(df, pkl_dir):
    """Suggest the target column based on historical data or column names."""
    log_step(
        "Suggesting target column", "To identify the appropriate target for modeling."
    )
    previous_data = load_pickle_data(pkl_dir)

    suggested_columns = set()

    for metadata in previous_data.values():
        if isinstance(metadata, dict) and "target_column" in metadata:
            suggested_columns.add(metadata["target_column"])

    for column in df.columns:
        if any(x in column.lower() for x in ["price", "target", "label"]):
            suggested_columns.add(column)

    suggested_target_column = next(iter(suggested_columns), None)
    if suggested_target_column:
        choice = (
            input(f"Use '{suggested_target_column}' as the target column? (y/n): ")
            .strip()
            .lower()
        )
        if choice == "y":
            return suggested_target_column

    print(f"Previous target columns: {suggested_columns}")
    prev_choice = input("Choose a previous target column? (y/n): ").strip().lower()
    if prev_choice == "y":
        prev_target_column = input(f"Choose from: {suggested_columns}\nYour choice: ")
        return prev_target_column

    target_column = input(f"Choose from: {list(df.columns)}\nYour choice: ")
    return target_column


def transform_data(df):
    """Impute missing values, normalize, and encode categorical data."""
    log_step(
        "Transforming data by imputing and normalizing",
        "To handle missing values and ensure data consistency.",
    )
    print("Starting data transformation...")

    if df.empty:
        raise ValueError("The DataFrame is empty. Please provide a valid dataset.")

    numerical_columns = df.select_dtypes(include=[np.number]).columns
    if len(numerical_columns) == 0:
        print("Only categorical data detected; skipping conversion.")
        return df, {}

    df_imputed = df.copy()
    imputer = SimpleImputer(strategy="mean")
    df_imputed[numerical_columns] = imputer.fit_transform(df_imputed[numerical_columns])

    if df_imputed[numerical_columns].max().max() != 0:
        normalizer = Normalizer()
        df_imputed[numerical_columns] = normalizer.fit_transform(
            df_imputed[numerical_columns]
        )

    label_encoders = {}
    categorical_columns = df.select_dtypes(include=["object"]).columns
    for column in categorical_columns:
        le = LabelEncoder()
        df_imputed[column] = le.fit_transform(df_imputed[column].astype(str))
        label_encoders[column] = le

    return df_imputed, label_encoders


def evaluate_models(df, target_column, plots_dir):
    """Evaluate different models based on the target column type."""
    log_step("Evaluating models", "To assess performance using suitable algorithms.")
    X = df.drop(columns=[target_column])
    y = df[target_column]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )

    models_info = {}
    for model in [
        ("Logistic Regression", LogisticRegression(max_iter=1000)),
        ("Random Forest Classifier", RandomForestClassifier()),
        ("KNN Classifier", KNeighborsClassifier()),
        ("Decision Tree Classifier", DecisionTreeClassifier()),
        ("Gaussian Naive Bayes", GaussianNB()),
        ("Gradient Boosting Classifier", GradientBoostingClassifier()),
    ]:
        model_name, model_instance = model
        model_instance.fit(X_train, y_train)
        y_pred = model_instance.predict(X_test)
        metrics = classification_report(
            y_test, y_pred, output_dict=True, zero_division=0
        )
        cm = confusion_matrix(y_test, y_pred)

        models_info[model_name] = {
            "metrics": metrics,
            "Confusion Matrix": cm,
            "features": X.columns.tolist(),
        }

        # Plotting confusion matrix
        fig = px.imshow(
            cm,
            text_auto=True,
            labels=dict(x="Predicted", y="True", color="Count"),
            title=f"Confusion Matrix: {model_name}",
        )
        plot_file = os.path.join(plots_dir, f"confusion_matrix_{model_name}.html")
        pio.write_html(fig, plot_file)

        # Save scatter plots of features vs target
        for feature in X.columns:
            sanitized_feature = sanitize_filename(feature)
            sanitized_target = sanitize_filename(target_column)
            scatter_fig = px.scatter(
                x=X_test[feature], y=y_test, title=f"{feature} vs {target_column}"
            )
            scatter_plot_file = os.path.join(
                plots_dir, f"scatter_{sanitized_feature}_vs_{sanitized_target}.html"
            )
            pio.write_html(scatter_fig, scatter_plot_file)

    return models_info


# Sanitize feature and target names for file paths
def sanitize_filename(name):
    return re.sub(r'[<>:"/\\|?*]', "_", name)


# Save scatter plots of features vs target
for feature in X.columns:
    sanitized_feature = sanitize_filename(feature)
    sanitized_target = sanitize_filename(target_column)
    scatter_fig = px.scatter(
        x=X_test[feature], y=y_test, title=f"{feature} vs {target_column}"
    )
    scatter_plot_file = os.path.join(
        plots_dir, f"scatter_{sanitized_feature}_vs_{sanitized_target}.html"
    )
    pio.write_html(scatter_fig, scatter_plot_file)


def main(dataset_path, output_directory, target_column=None):
    """Main function to execute the data pipeline."""
    setup_logging("data_analysis.log")

    log_step("Reading dataset", "Load dataset from the provided path.")
    df = pd.read_csv(dataset_path)

    created_dirs = create_directories(output_directory, os.path.basename(dataset_path))

    log_step("Checking column names", "Ensure column names are standardized.")
    df = check_column_names(df)

    if target_column is None:
        log_step(
            "Suggesting target column",
            "To identify the appropriate target for modeling.",
        )
        target_column = suggest_target_column(df, created_dirs["analysis"])

    log_step("Transforming dataset", "Prepare the data for analysis.")
    df_transformed, label_encoders = transform_data(df)

    log_step("Evaluating models", "Assess performance using various algorithms.")
    models_info = evaluate_models(df_transformed, target_column, created_dirs["plots"])

    best_model = max(
        models_info.items(), key=lambda x: x[1]["metrics"]["accuracy"], default=None
    )

    save_analysis_log(
        analysis_log,
        created_dirs["analysis"],
        models_info,
        best_model[0] if best_model else None,
        created_dirs["plots"],
        target_column,
    )

    save_to_pickle(
        df,
        created_dirs["analysis"],
        target_column,
        best_model[0] if best_model else None,
        label_encoders,
        None,
        os.path.basename(dataset_path),
    )


if __name__ == "__main__":
    dataset_path = r"C:\Users\HimakarRaju\Desktop\Milestone2\Datasets2\bsf.csv"  # Specify your dataset path here
    output_directory = "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts2"  # Specify your output directory here
    main(dataset_path, output_directory)
