import os
import pickle

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from matplotlib.backends.backend_pdf import PdfPages
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    confusion_matrix,
)
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.svm import SVC
from sklearn.tree import DecisionTreeClassifier


# Function to read the file (CSV or Excel)
def load_file(file_path):
    """Loads a CSV or Excel file into a DataFrame."""
    if file_path.endswith(".csv"):
        df = pd.read_csv(file_path)
    elif file_path.endswith(".xlsx"):
        df = pd.read_excel(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a CSV or Excel file.")
    return df


# Function to perform data transformation
def transform_data(df):
    """Transform the dataset for model usage."""
    # Handling missing values
    imputer = SimpleImputer(strategy="mean")
    df_imputed = pd.DataFrame(
        imputer.fit_transform(df.select_dtypes(include=[np.number])),
        columns=df.select_dtypes(include=[np.number]).columns,
    )

    # Label Encoding for categorical columns
    label_encoders = {}
    for column in df.select_dtypes(include=["object"]).columns:
        le = LabelEncoder()
        df[column] = le.fit_transform(df[column].astype(str))
        label_encoders[column] = le

    # Scaling numerical features
    scaler = StandardScaler()
    df[df_imputed.columns] = scaler.fit_transform(df_imputed)

    return df, label_encoders, scaler


# Save transformation methods to Word document
def save_transformation_methods(df, doc_filename):
    """Saves the transformation methods used to a Word document."""
    doc = Document()
    doc.add_heading("Data Transformation Report", 0)

    doc.add_heading("1. Missing Data Imputation:", level=1)
    doc.add_paragraph(
        "Missing values were imputed using the SimpleImputer with mean strategy."
    )

    doc.add_heading("2. Label Encoding for Categorical Columns:", level=1)
    doc.add_paragraph(
        "Label Encoding was used to transform categorical variables into numerical values."
    )

    doc.add_heading("3. Feature Scaling:", level=1)
    doc.add_paragraph(
        "StandardScaler was used to scale numerical features to have a mean of 0 and variance of 1."
    )

    doc.add_heading("Transformed Data Preview:", level=1)
    doc.add_paragraph(str(df.head()))

    doc.save(doc_filename)


# Function to perform scikit-learn methods
def scikit_analysis(df, target_column):
    """Perform scikit-learn methods for model selection and analysis."""
    X = df.drop(columns=[target_column])
    y = df[target_column]

    # Split the data into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )

    # Define a list of models
    models = {
        "Logistic Regression": LogisticRegression(),
        "Random Forest": RandomForestClassifier(),
        "Support Vector Classifier": SVC(),
        "Decision Tree": DecisionTreeClassifier(),
        "K-Nearest Neighbors": KNeighborsClassifier(),
        "Naive Bayes": GaussianNB(),
        "Gradient Boosting": GradientBoostingClassifier(),
    }

    # Cross-validation scores and model evaluation
    results = {}
    best_model_name = None
    best_model_score = 0

    for model_name, model in models.items():
        model.fit(X_train, y_train)
        y_pred = model.predict(X_test)

        accuracy = accuracy_score(y_test, y_pred)
        precision = precision_score(y_test, y_pred, average="weighted", zero_division=0)
        recall = recall_score(y_test, y_pred, average="weighted", zero_division=0)
        f1 = f1_score(y_test, y_pred, average="weighted", zero_division=0)
        conf_matrix = confusion_matrix(y_test, y_pred)

        # Storing model metrics
        results[model_name] = {
            "accuracy": accuracy,
            "precision": precision,
            "recall": recall,
            "f1_score": f1,
            "confusion_matrix": conf_matrix,
        }

        # Keep track of the best model
        if accuracy > best_model_score:
            best_model_score = accuracy
            best_model_name = model_name

    return results, best_model_name, X_train, X_test, y_train, y_test


# Save analysis and explanation to Word document
def save_analysis_to_word(results, best_model_name, doc_filename):
    """Saves analysis results and reasoning for model selection to a Word document."""
    doc = Document()
    doc.add_heading("Model Evaluation Report", 0)

    doc.add_heading("Model Comparison Results:", level=1)

    # Write metrics for all models
    for model_name, metrics in results.items():
        doc.add_heading(f"{model_name} Results:", level=2)
        doc.add_paragraph(f"Accuracy: {metrics['accuracy']:.4f}")
        doc.add_paragraph(f"Precision: {metrics['precision']:.4f}")
        doc.add_paragraph(f"Recall: {metrics['recall']:.4f}")
        doc.add_paragraph(f"F1 Score: {metrics['f1_score']:.4f}")
        doc.add_paragraph(f"Confusion Matrix: \n{metrics['confusion_matrix']}")

    # Write reasoning for choosing the best model
    doc.add_heading("Best Model Selection:", level=1)
    doc.add_paragraph(
        f"The best model chosen is {best_model_name}. "
        f"The reason for this selection is based on its highest accuracy score compared to other models."
    )

    doc.save(doc_filename)


# Plot graphs and save to PDF
def plot_graphs(df, file_name):
    """Generate plots and save them as a PDF."""
    # Numerical columns
    numerical_cols = df.select_dtypes(include=[np.number]).columns

    # Create a PDF file to save plots
    with PdfPages(file_name) as pdf:
        for col in numerical_cols:
            plt.figure(figsize=(10, 6))
            sns.histplot(df[col], kde=True)
            plt.title(f"Distribution of {col}")
            pdf.savefig()
            plt.close()

        # Plot correlation matrix
        plt.figure(figsize=(12, 8))
        sns.heatmap(df.corr(), annot=True, cmap="coolwarm")
        plt.title("Correlation Matrix")
        pdf.savefig()
        plt.close()


# Save the model
def save_model(model, filename):
    """Save the trained model to a pickle file."""
    with open(filename, "wb") as f:
        pickle.dump(model, f)


# Function to process all files in a directory
def process_files(input_path, output_dir):
    """Process all files in a directory or a single file."""
    if os.path.isdir(input_path):
        # Process all CSV files in the directory
        for file_name in os.listdir(input_path):
            if file_name.endswith(".csv"):
                file_path = os.path.join(input_path, file_name)
                print(f"Processing file: {file_path}")
                process_single_file(file_path, output_dir)
    elif os.path.isfile(input_path) and input_path.endswith(".csv"):
        # Process a single file
        process_single_file(input_path, output_dir)
    else:
        print("Invalid file or directory path.")


# Function to process a single file
def process_single_file(file_path, output_dir):
    """Process a single CSV file."""
    # Load the data
    df = load_file(file_path)

    # Automatically choose the last column as the target column for now
    target_column = df.columns[-1]
    print(f"Automatically selected target column: {target_column}")

    # Create respective directories for storing results
    base_filename = os.path.splitext(os.path.basename(file_path))[0]
    csv_folder = os.path.join(output_dir, "csv_files")
    plots_folder = os.path.join(output_dir, "plots")
    models_folder = os.path.join(output_dir, "models")
    analysis_folder = os.path.join(output_dir, "analysis")

    os.makedirs(csv_folder, exist_ok=True)
    os.makedirs(plots_folder, exist_ok=True)
    os.makedirs(models_folder, exist_ok=True)
    os.makedirs(analysis_folder, exist_ok=True)

    # Step 1: Transform the data
    transformed_df, label_encoders, scaler = transform_data(df)

    # Step 2: Save transformation methods
    save_transformation_methods(
        transformed_df,
        os.path.join(analysis_folder, f"transformation_methods_{base_filename}.docx"),
    )

    # Step 3: Perform Scikit analysis and evaluate models
    results, best_model_name, X_train, X_test, y_train, y_test = scikit_analysis(
        transformed_df, target_column
    )

    # Step 4: Save analysis report
    save_analysis_to_word(
        results,
        best_model_name,
        os.path.join(analysis_folder, f"analysis_{base_filename}.docx"),
    )

    # Step 5: Plot and save figures
    plot_graphs(
        transformed_df, os.path.join(plots_folder, f"plots_{base_filename}.pdf")
    )

    # Step 6: Save the best model
    best_model = results[best_model_name]
    save_model(best_model, os.path.join(models_folder, f"model_{base_filename}.pkl"))

    # Step 7: Save the transformed data
    transformed_df.to_csv(
        os.path.join(csv_folder, f"transformed_{base_filename}.csv"), index=False
    )

    print(f"Processing completed for {file_path}")

    # Step 8: Ask user if they want to change the target column
    print("Available columns for target: ")
    available_columns = df.columns.tolist()
    for idx, col in enumerate(available_columns, 1):
        print(f"{idx}. {col}")

    # Automatically select the last column as the initial target column
    target_column = available_columns[-1]
    print(f"Automatically selected target column: {target_column}")

    change_target = (
        input("Do you want to change the target column? (yes/no): ").strip().lower()
    )

    if change_target == "yes":
        new_target_index = (
            int(input(f"Choose a new target column (1-{len(available_columns)}): ")) - 1
        )
        if 0 <= new_target_index < len(available_columns):
            new_target_column = available_columns[new_target_index]
            print(f"Rerunning the analysis with target column: {new_target_column}")
            # Here, you would call your function to reprocess with the new target column
        else:
            print(f"Error: Invalid selection.")
    else:
        print("Target column confirmed.")


# Main function to execute the script
if __name__ == "__main__":
    input_file = input(
        "Enter path for file/folder : "
    )  # Single file or folder containing files
    output_folder = input(
        "Enter path for Data Read Outs : "
    )  # Directory to store results

    process_files(input_file, output_folder)
