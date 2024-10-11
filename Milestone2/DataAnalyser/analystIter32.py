import os

import numpy as np
import pandas as pd
import plotly.express as px
from docx import Document
from sklearn.decomposition import PCA
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    RandomForestRegressor,
    GradientBoostingRegressor,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    confusion_matrix,
    mean_squared_error,
    mean_absolute_error,
    r2_score,
)
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.svm import SVC, SVR
from sklearn.tree import DecisionTreeClassifier


# Function to create directories
def create_directories(output_directory, dataset_name):
    """Create directories for analysis reports and plots."""
    analysis_dir = os.path.join(output_directory, dataset_name, "analysis")
    plots_dir = os.path.join(output_directory, dataset_name, "plots")
    os.makedirs(analysis_dir, exist_ok=True)
    os.makedirs(plots_dir, exist_ok=True)
    return analysis_dir, plots_dir


# Function to save plots as HTML
def save_plot_as_html(plot, plot_name, plots_dir):
    """Save plot as an interactive HTML file."""
    plot_path = os.path.join(plots_dir, f"{plot_name}.html")
    plot.write_html(plot_path)


# Function to read the file (CSV or Excel)
def load_file(file_path):
    """Loads a CSV or Excel file into a DataFrame."""
    try:
        if file_path.endswith(".csv"):
            df = pd.read_csv(file_path)
        elif file_path.endswith(".xlsx"):
            df = pd.read_excel(file_path)
        else:
            raise ValueError(
                "Unsupported file format. Please provide a CSV or Excel file."
            )
    except Exception as e:
        print(f"Error loading file {file_path}: {e}")
        df = pd.DataFrame()  # Return empty DataFrame if file load fails
    return df


# Function to suggest target column using NLP heuristics
def suggest_target_column(df):
    """Suggest the target column using simple heuristics and NLP."""
    target_column = None

    # First: Check for columns that contain words like "price", "target", or "label"
    for column in df.columns:
        column_name = column.lower()
        if "price" in column_name or "target" in column_name or "label" in column_name:
            target_column = column
            break

    # Second: Check if there's a column with fewer unique values for classification tasks
    if not target_column:
        for column in df.columns:
            if len(df[column].unique()) <= 10:  # Heuristic for classification
                target_column = column
                break

    # If no suggestion, ask the user manually
    if not target_column:
        target_column = input(
            "Could not automatically detect the target column. Please enter the target column name: "
        )

    return target_column


# Function to perform data transformation
def transform_data(df):
    """Transform the dataset for model usage."""
    transformation_report = []

    # Handling missing values using SimpleImputer (mean strategy for numerical columns)
    imputer = SimpleImputer(strategy="mean")
    df_imputed = df.copy()  # To avoid SettingWithCopyWarning
    numerical_columns = df.select_dtypes(include=[np.number]).columns
    df_imputed[numerical_columns] = imputer.fit_transform(df_imputed[numerical_columns])
    transformation_report.append(
        "Missing values were imputed using the SimpleImputer with mean strategy."
    )

    # Check if normalization is necessary, avoid if it leads to 0 values
    normalizer = Normalizer()
    if (
        df_imputed[numerical_columns].max().max() != 0
    ):  # Check if normalization is needed
        df_imputed[numerical_columns] = normalizer.fit_transform(
            df_imputed[numerical_columns]
        )
        transformation_report.append("Normalization was applied to numerical features.")
    else:
        transformation_report.append(
            "Normalization skipped due to zero values in data."
        )

    # Label Encoding for categorical columns
    label_encoders = {}
    categorical_columns = df.select_dtypes(include=["object"]).columns
    for column in categorical_columns:
        le = LabelEncoder()
        df_imputed[column] = le.fit_transform(df[column].astype(str))
        label_encoders[column] = le
    transformation_report.append("Label Encoding was applied to categorical variables.")

    return df_imputed, label_encoders, transformation_report


# Apply PCA to the transformed data
def apply_pca(X, n_components=0.95):
    """Apply PCA for dimensionality reduction."""
    try:
        pca = PCA(n_components=n_components)
        X_pca = pca.fit_transform(X)
        return X_pca, pca
    except Exception as e:
        print(f"Error during PCA: {e}")
        return X, None


# Generate useful plots and save as HTML
def generate_plots(df, target_column, plots_dir):
    """Generate useful plots (feature distributions, PCA, correlation heatmap)."""
    # Feature distribution plot (for numerical columns)
    for column in df.select_dtypes(include=[np.number]).columns:
        fig = px.histogram(df, x=column, title=f"Distribution of {column}")
        save_plot_as_html(fig, f"feature_distribution_{column}", plots_dir)

    # Correlation heatmap
    corr_matrix = df.corr()
    fig = px.imshow(corr_matrix, title="Correlation Heatmap")
    save_plot_as_html(fig, "correlation_heatmap", plots_dir)

    # PCA plot (for dimensionality reduction)
    X = df.drop(columns=[target_column])
    X_pca, _ = apply_pca(X, n_components=2)
    pca_df = pd.DataFrame(X_pca, columns=["PCA1", "PCA2"])
    pca_df[target_column] = df[target_column]
    fig = px.scatter(pca_df, x="PCA1", y="PCA2", color=target_column, title="PCA Plot")
    save_plot_as_html(fig, "pca_plot", plots_dir)


# Save transformation methods to Word document
def save_transformation_methods(df, doc_filename, transformation_report):
    """Saves the transformation methods used to a Word document."""
    doc = Document()
    doc.add_heading("Data Transformation Report", 0)

    doc.add_heading("1. Missing Data Imputation:", level=1)
    doc.add_paragraph(transformation_report[0])

    doc.add_heading("2. Label Encoding for Categorical Columns:", level=1)
    doc.add_paragraph(transformation_report[2])

    doc.add_heading("3. Feature Normalization:", level=1)
    doc.add_paragraph(transformation_report[1])

    doc.add_heading("Transformed Data Preview:", level=1)
    doc.add_paragraph(str(df.head()))

    doc.save(doc_filename)


# Function to save analysis to a Word document
def save_analysis_to_word(
    results, doc_filename, is_classification, df, transformation_report
):
    """Saves the analysis results to a Word document."""
    doc = Document()
    doc.add_heading("Model Analysis Report", 0)

    # Transformation section
    doc.add_heading("Data Transformation Steps:", level=1)
    doc.add_paragraph("\n".join(transformation_report))

    # Model evaluation section
    if is_classification:
        doc.add_heading("Classification Results", level=1)
        for model_name, metrics in results.items():
            doc.add_heading(f"Model: {model_name}", level=2)
            doc.add_paragraph(f"Accuracy: {metrics['accuracy']}")
            doc.add_paragraph(f"Precision: {metrics['precision']}")
            doc.add_paragraph(f"Recall: {metrics['recall']}")
            doc.add_paragraph(f"F1 Score: {metrics['f1_score']}")
            doc.add_paragraph(f"Confusion Matrix: {metrics['confusion_matrix']}")
    else:
        doc.add_heading("Regression Results", level=1)
        for model_name, metrics in results.items():
            doc.add_heading(f"Model: {model_name}", level=2)
            doc.add_paragraph(f"Mean Squared Error: {metrics['mse']}")
            doc.add_paragraph(f"Mean Absolute Error: {metrics['mae']}")
            doc.add_paragraph(f"R2 Score: {metrics['r2']}")

    doc.save(doc_filename)


# Main function to invoke all processes
def run_analysis(file_path, output_directory):
    """Runs the entire pipeline of loading, transforming, analyzing, and saving reports."""
    try:
        df = load_file(file_path)
        if df.empty:
            print("No data to process.")
            return

        # Suggest the target column and transform the data
        target_column = suggest_target_column(df)
        print(f"target_column : {target_column}")

        print("Transforming Data")
        df_imputed, label_encoders, transformation_report = transform_data(df)

        # Create directories
        dataset_name = os.path.basename(file_path).split(".")[0]
        analysis_dir, plots_dir = create_directories(output_directory, dataset_name)

        # Split the data into training and testing sets
        X = df_imputed.drop(columns=[target_column])
        y = df_imputed[target_column]
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=0.3, random_state=42
        )

        # Apply PCA for dimensionality reduction
        X_train_pca, pca = apply_pca(X_train, n_components=0.95)
        X_test_pca = pca.transform(X_test) if pca else X_test

        # Check if the target is a classification or regression problem
        is_classification = (
            len(np.unique(y)) < 20
        )  # Heuristic: fewer unique values for classification

        # Models for analysis
        models = {}
        results = {}

        if is_classification:
            models = {
                "Logistic Regression": LogisticRegression(max_iter=200),
                "Random Forest": RandomForestClassifier(),
                "Support Vector Classifier": SVC(),
                "Decision Tree": DecisionTreeClassifier(),
                "K-Nearest Neighbors": KNeighborsClassifier(),
                "Naive Bayes": GaussianNB(),
                "Gradient Boosting": GradientBoostingClassifier(),
            }

            # Evaluate models
            for model_name, model in models.items():
                try:
                    model.fit(X_train_pca, y_train)
                    y_pred = model.predict(X_test_pca)

                    accuracy = accuracy_score(y_test, y_pred)
                    precision = precision_score(
                        y_test, y_pred, average="weighted", zero_division=0
                    )
                    recall = recall_score(
                        y_test, y_pred, average="weighted", zero_division=0
                    )
                    f1 = f1_score(y_test, y_pred, average="weighted", zero_division=0)
                    cm = confusion_matrix(y_test, y_pred)

                    results[model_name] = {
                        "accuracy": accuracy,
                        "precision": precision,
                        "recall": recall,
                        "f1_score": f1,
                        "confusion_matrix": cm,
                    }

                except Exception as e:
                    print(f"Error training model {model_name}: {e}")

            # Save analysis to Word
            save_analysis_to_word(
                results,
                os.path.join(analysis_dir, "model_analysis.docx"),
                is_classification,
                df_imputed,
                transformation_report,
            )

        else:
            models = {
                "Linear Regression": LinearRegression(),
                "Random Forest Regressor": RandomForestRegressor(),
                "SVR": SVR(),
                "K-Nearest Neighbors Regressor": KNeighborsRegressor(),
                "Gradient Boosting Regressor": GradientBoostingRegressor(),
            }

            # Evaluate regression models
            for model_name, model in models.items():
                try:
                    model.fit(X_train_pca, y_train)
                    y_pred = model.predict(X_test_pca)

                    mse = mean_squared_error(y_test, y_pred)
                    mae = mean_absolute_error(y_test, y_pred)
                    r2 = r2_score(y_test, y_pred)

                    results[model_name] = {
                        "mse": mse,
                        "mae": mae,
                        "r2": r2,
                    }

                except Exception as e:
                    print(f"Error training model {model_name}: {e}")

            # Save regression analysis to Word
            save_analysis_to_word(
                results,
                os.path.join(analysis_dir, "regression_analysis.docx"),
                is_classification,
                df_imputed,
                transformation_report,
            )

        # Generate and save plots
        generate_plots(df_imputed, target_column, plots_dir)

        # Save transformation report
        save_transformation_methods(
            df_imputed,
            os.path.join(analysis_dir, "transformation_report.docx"),
            transformation_report,
        )

        print("Analysis completed and saved.")
    except Exception as e:
        print(f"Error during analysis: {e}")


if __name__ == "__main__":
    run_analysis(
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\Datasets\\gym_members_exercise_tracking.csv",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts",
    )
