import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.svm import SVC
from sklearn.tree import DecisionTreeClassifier
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import GaussianNB
from sklearn.model_selection import cross_val_score
from docx import Document
import pickle

# Function to read the file (CSV or Excel)
def load_file(file_path):
    """Loads a CSV or Excel file into a DataFrame."""
    if file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    elif file_path.endswith('.xlsx'):
        df = pd.read_excel(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a CSV or Excel file.")
    return df

# Function to perform data transformation
def transform_data(df):
    """Transform the dataset for model usage."""
    # Handling missing values
    imputer = SimpleImputer(strategy='mean')
    df_imputed = pd.DataFrame(imputer.fit_transform(df.select_dtypes(include=[np.number])), columns=df.select_dtypes(include=[np.number]).columns)
    
    # Label Encoding for categorical columns
    label_encoders = {}
    for column in df.select_dtypes(include=['object']).columns:
        le = LabelEncoder()
        df[column] = le.fit_transform(df[column].astype(str))
        label_encoders[column] = le
    
    # Scaling numerical features
    scaler = StandardScaler()
    df[df_imputed.columns] = scaler.fit_transform(df_imputed)
    
    return df, label_encoders, scaler

# Save transformation methods to Word document
def save_transformation_methods(df, doc_filename):
    """Saves the transformation methods used to a Word document."""
    doc = Document()
    doc.add_heading('Data Transformation Report', 0)
    
    doc.add_heading('1. Missing Data Imputation:', level=1)
    doc.add_paragraph('Missing values were imputed using the SimpleImputer with mean strategy.')
    
    doc.add_heading('2. Label Encoding for Categorical Columns:', level=1)
    doc.add_paragraph('Label Encoding was used to transform categorical variables into numerical values.')
    
    doc.add_heading('3. Feature Scaling:', level=1)
    doc.add_paragraph('StandardScaler was used to scale numerical features to have a mean of 0 and variance of 1.')
    
    doc.add_heading('Transformed Data Preview:', level=1)
    doc.add_paragraph(str(df.head()))
    
    doc.save(doc_filename)

# Function to perform scikit-learn methods
def scikit_analysis(df, target_column):
    """Perform scikit-learn methods for model selection and analysis."""
    X = df.drop(columns=[target_column])
    y = df[target_column]
    
    # Split the data into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    # Define a list of models
    models = {
        'Logistic Regression': LogisticRegression(),
        'Random Forest': RandomForestClassifier(),
        'Support Vector Classifier': SVC(),
        'Decision Tree': DecisionTreeClassifier(),
        'K-Nearest Neighbors': KNeighborsClassifier(),
        'Naive Bayes': GaussianNB(),
        'Gradient Boosting': GradientBoostingClassifier()
    }
    
    # Cross-validation scores and model evaluation
    results = {}
    best_model_name = None
    best_model_score = 0
    
    for model_name, model in models.items():
        model.fit(X_train, y_train)
        y_pred = model.predict(X_test)
        
        accuracy = accuracy_score(y_test, y_pred)
        precision = precision_score(y_test, y_pred, average='weighted', zero_division=0)
        recall = recall_score(y_test, y_pred, average='weighted', zero_division=0)
        f1 = f1_score(y_test, y_pred, average='weighted', zero_division=0)
        conf_matrix = confusion_matrix(y_test, y_pred)
        
        # Storing model metrics
        results[model_name] = {
            'accuracy': accuracy,
            'precision': precision,
            'recall': recall,
            'f1_score': f1,
            'confusion_matrix': conf_matrix
        }
        
        # Keep track of the best model
        if accuracy > best_model_score:
            best_model_score = accuracy
            best_model_name = model_name
    
    return results, best_model_name, X_train, X_test, y_train, y_test

# Save analysis and explanation to Word document
def save_analysis_to_word(results, best_model_name, doc_filename):
    """Saves analysis results and reasoning for model selection to a Word document."""
    doc = Document()
    doc.add_heading('Model Evaluation Report', 0)
    
    doc.add_heading('Model Comparison Results:', level=1)
    
    # Write metrics for all models
    for model_name, metrics in results.items():
        doc.add_heading(f'{model_name} Results:', level=2)
        doc.add_paragraph(f"Accuracy: {metrics['accuracy']:.4f}")
        doc.add_paragraph(f"Precision: {metrics['precision']:.4f}")
        doc.add_paragraph(f"Recall: {metrics['recall']:.4f}")
        doc.add_paragraph(f"F1 Score: {metrics['f1_score']:.4f}")
        doc.add_paragraph(f"Confusion Matrix: \n{metrics['confusion_matrix']}")
    
    # Write reasoning for choosing the best model
    doc.add_heading('Best Model Selection:', level=1)
    doc.add_paragraph(f"The best model chosen is {best_model_name}. The reason for this selection is based on its highest accuracy score compared to other models.")
    
    doc.save(doc_filename)

# Plot graphs and save to PDF
def plot_graphs(df, file_name):
    """Generate plots and save them as a PDF."""
    # Numerical columns
    numerical_cols = df.select_dtypes(include=[np.number]).columns
    
    # Create a PDF file to save plots
    from matplotlib.backends.backend_pdf import PdfPages
    with PdfPages(file_name) as pdf:
        for col in numerical_cols:
            plt.figure(figsize=(10, 6))
            sns.histplot(df[col], kde=True)
            plt.title(f'Distribution of {col}')
            pdf.savefig()
            plt.close()

        # Plot correlation matrix
        plt.figure(figsize=(12, 8))
        sns.heatmap(df.corr(), annot=True, cmap='coolwarm')
        plt.title('Correlation Matrix')
        pdf.savefig()
        plt.close()

# Save the model
def save_model(model, filename):
    """Save the trained model to a pickle file."""
    with open(filename, 'wb') as f:
        pickle.dump(model, f)

# Main function to execute the workflow
def main(file_path, target_column):
    """Main function to load, transform, analyze, and save the results."""
    # Step 1: Load the data
    df = load_file(file_path)
    
    # Step 2: Transform the data
    transformed_df, label_encoders, scaler = transform_data(df)
    
    # Step 3: Save transformation methods
    save_transformation_methods(transformed_df, f"transformation_methods_{file_path.split('/')[-1].split('.')[0]}.docx")
    
    # Step 4: Perform scikit-learn analysis
    results, best_model_name, X_train, X_test, y_train, y_test = scikit_analysis(transformed_df, target_column)
    
    # Step 5: Save analysis report
    save_analysis_to_word(results, best_model_name, f"analysis_on_{file_path.split('/')[-1].split('.')[0]}.docx")
    
    # Step 6: Plot results and save to PDF
    plot_graphs(transformed_df, f"plots_{file_path.split('/')[-1].split('.')[0]}.pdf")
    
    # Step 7: Save transformed data to CSV
    transformed_df.to_csv(f"transformed_{file_path.split('/')[-1]}", index=False)
    
    # Step 8: Save the best model
    best_model = models[best_model_name]
    best_model.fit(X_train, y_train)
    save_model(best_model, f"model_{file_path.split('/')[-1].split('.')[0]}.pkl")
    
    # Step 9: Evaluate the model
    y_pred = best_model.predict(X_test)
    print("Model Evaluation Metrics:")
    print(classification_report(y_test, y_pred))
    print("Confusion Matrix:")
    print(confusion_matrix(y_test, y_pred))

# Execute the program
if __name__ == "__main__":
    file_path = input("Enter path for file/folder")  # Change this to the actual file path
    target_column = "target"  # Change this to the actual target column name
    main(file_path, target_column)
