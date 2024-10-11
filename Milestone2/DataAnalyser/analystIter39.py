import logging
import os
import pickle
import re

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
from docx import Document
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import LabelEncoder, Normalizer

# Setup Plotly renderer
pio.renderers.default = "browser"


# Setup logging
def setup_logging(log_file):
    logging.basicConfig(
        filename=log_file,
        level=logging.ERROR,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


# Check for symbols and spaces in column names
def check_column_names(df):
    invalid_columns = [col for col in df.columns if re.search(r"[^a-zA-Z0-9_]", col)]
    if invalid_columns:
        print(f"Warning: Columns with symbols/spaces detected: {invalid_columns}")
        df.columns = df.columns.str.replace(r"[^a-zA-Z0-9_]", "_", regex=True)
    return df


# Create directories
def create_directories(output_directory, dataset_name):
    dirs = ["analysis", "plots", "logs", "pkl", "time_series"]
    return {d: os.path.join(output_directory, dataset_name, d) for d in dirs}


# Load target column data from pickle
def load_target_from_pickle(pkl_dir):
    pkl_file = os.path.join(pkl_dir, "target_column.pkl")
    if os.path.exists(pkl_file):
        with open(pkl_file, "rb") as f:
            target_column = pickle.load(f)
        print(f"Previous target column found: {target_column}")
        return target_column
    return None


# Save target column to pickle
def save_target_to_pickle(target_column, pkl_dir):
    pkl_file = os.path.join(pkl_dir, "target_column.pkl")
    with open(pkl_file, "wb") as f:
        pickle.dump(target_column, f)


# Suggest target column
def suggest_target_column(df, pkl_dir):
    target_column = load_target_from_pickle(pkl_dir)
    if not target_column:
        for column in df.columns:
            if any(x in column.lower() for x in ["price", "target", "label"]):
                target_column = column
                break
        if not target_column:
            target_column = input(
                f"No suitable target column found. Choose from: {list(df.columns)} "
            )

    print(f"Suggested target column: {target_column}")
    confirmation = input("Is this okay? (y/n): ")
    if confirmation.lower() != "y":
        target_column = input(
            f"Please choose the target column from: {list(df.columns)} "
        )

    save_target_to_pickle(target_column, pkl_dir)
    return target_column


# Data transformation
def transform_data(df):
    print("Starting data transformation...")
    imputer = SimpleImputer(strategy="mean")
    df_imputed = df.copy()

    numerical_columns = df.select_dtypes(include=[np.number]).columns
    df_imputed[numerical_columns] = imputer.fit_transform(df_imputed[numerical_columns])

    # Normalization
    if df_imputed[numerical_columns].max().max() != 0:
        normalizer = Normalizer()
        df_imputed[numerical_columns] = normalizer.fit_transform(
            df_imputed[numerical_columns]
        )

    # Label encoding for categorical variables
    label_encoders = {}
    categorical_columns = df.select_dtypes(include=["object"]).columns
    for column in categorical_columns:
        le = LabelEncoder()
        df_imputed[column] = le.fit_transform(df_imputed[column].astype(str))
        label_encoders[column] = le

    print("Data transformation completed.")
    return df_imputed, label_encoders


# Log errors
def log_error(model_name, error, log_file):
    logging.error(f"Model {model_name} failed with error: {error}")


# Save the best model and its metrics
def save_best_model(results, doc_filename):
    if results:
        best_model = max(results, key=lambda model: results[model]["r2"])
        print(f"Best Model: {best_model} with R2: {results[best_model]['r2']}")
        with open(doc_filename, "a") as f:
            f.write(f"Best Model: {best_model} - Metrics: {results[best_model]}\n")
    else:
        print("No models were successfully trained.")


# Generate and save plots
def generate_and_save_plots(df, plots_dir, target_column):
    print("Generating plots...")

    # Histograms for numerical features with KDE
    numerical_cols = df.select_dtypes(include=[np.number]).columns
    for col in numerical_cols:
        fig = px.histogram(
            df,
            x=col,
            title=f"Histogram of {col}",
            marginal="box",
            histnorm="probability density",
        )
        fig.write_html(os.path.join(plots_dir, f"histogram_{col}.html"))

    # Correlation matrix
    corr = df.corr()
    fig = px.imshow(corr, text_auto=True, title="Correlation Matrix")
    fig.write_html(os.path.join(plots_dir, "correlation_matrix.html"))

    # Target vs useful features with KDE
    for col in numerical_cols:
        if col != target_column:
            fig = px.scatter(
                df,
                x=col,
                y=target_column,
                title=f"{target_column} vs {col}",
                trendline="lowess",
            )  # Add a trendline for KDE
            fig.write_html(os.path.join(plots_dir, f"{target_column}_vs_{col}.html"))


# Time series analysis
def parse_dates(df, date_columns):
    # Define potential date formats
    date_formats = [
        "%Y-%m-%d",  # e.g., 2024-10-11
        "%m/%d/%Y",  # e.g., 10/11/2024
        "%d-%b-%Y",  # e.g., 11-Oct-2024
        "%B %d, %Y",  # e.g., October 11, 2024
        "%Y%m%d",  # e.g., 20241011
        "%d/%m/%Y",  # e.g., 11/10/2024
        "%Y-%m",  # e.g., 2024-10
        "%Y",  # e.g., 2024
    ]

    for col in date_columns:
        for fmt in date_formats:
            try:
                df[col] = pd.to_datetime(df[col], format=fmt, errors="coerce")
                if df[col].isnull().sum() == 0:  # If all values are valid
                    print(f"Successfully parsed {col} with format {fmt}.")
                    break
            except Exception as e:
                print(
                    f"Could not parse {col} with format {fmt}. Continuing to next format."
                )

    return df


def perform_time_series_analysis(df, time_series_dir):
    date_columns = df.select_dtypes(include=["datetime"]).columns
    if date_columns.empty:
        print("No date columns found for time series analysis.")
        return

    for col in date_columns:
        try:
            numerical_col = df.select_dtypes(include=[np.number]).columns[
                0
            ]  # Use the first numerical column for plotting
            # Drop NaNs from the time series analysis
            filtered_df = df[[col, numerical_col]].dropna()
            fig = px.line(
                filtered_df,
                x=col,
                y=numerical_col,
                title=f"Time Series Analysis for {col}",
            )
            fig.write_html(os.path.join(time_series_dir, f"time_series_{col}.html"))
            print(f"Time series plot saved for {col}.")
        except Exception as e:
            log_error(col, e, "time_series_error.log")


# Document analysis in a Word file
def create_analysis_document(
    analysis_dir, df, results, plots_dir, transformation_report
):
    doc = Document()

    # Title
    doc.add_heading("Data Analysis Report", level=1)

    # Introduction
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This document outlines the analysis performed on the dataset, including model evaluations, visualizations, and insights."
    )

    # Data Overview
    doc.add_heading("Data Overview", level=2)
    doc.add_paragraph(f"Dataset Shape: {df.shape}")
    doc.add_paragraph(f"Columns: {list(df.columns)}")
    doc.add_paragraph("First few rows of the dataset:")
    doc.add_paragraph(str(df.head()))

    # Data Transformation Steps
    doc.add_heading("Data Transformation Steps", level=2)
    for step in transformation_report:
        doc.add_paragraph(step)

    # Models and Results
    doc.add_heading("Model Evaluations", level=2)
    for model_name, metrics in results.items():
        doc.add_heading(model_name, level=3)
        for key, value in metrics.items():
            doc.add_paragraph(f"{key}: {value}")

    # Plots Section
    doc.add_heading("Plots", level=2)
    for plot in os.listdir(plots_dir):
        doc.add_paragraph(plot)
        doc.add_picture(os.path.join(plots_dir, plot))

    # Insights and Conclusions
    doc.add_heading("Insights and Conclusions", level=2)
    doc.add_paragraph("Based on the analysis, the following insights were derived:")
    # Here you can add specific insights based on your analysis results
    doc.add_paragraph("1. Insert specific insights here.")
    doc.add_paragraph(
        "2. Additional insights can be added based on model performance or data trends."
    )

    # Save document
    doc_file = os.path.join(analysis_dir, "Data_Analysis_Report.docx")
    doc.save(doc_file)
    print(f"Analysis report saved as {doc_file}.")


def save_all_data_to_pickle(df, pkl_dir):
    pkl_file = os.path.join(pkl_dir, "full_data.pkl")
    with open(pkl_file, "wb") as f:
        pickle.dump(df, f)
    print("All data saved to pickle file.")


# Main analysis function
def run_analysis(file_path, output_directory):
    try:
        df = load_file(file_path)
        if df.empty:
            print("No data to process.")
            return

        # Create directories
        dataset_name = os.path.basename(file_path).split(".")[0]
        analysis_dir, plots_dir, logs_dir, pkl_dir = create_directories(
            output_directory, dataset_name
        )

        # Setup logging
        log_file = os.path.join(logs_dir, f"{dataset_name}_error_log.log")
        setup_logging(log_file)

        # Check column names and parse dates
        invalid_columns = check_column_names(df)
        if invalid_columns:
            df.columns = df.columns.str.replace(r"[^a-zA-Z0-9_]", "_")

        date_columns = df.select_dtypes(include=["object"]).columns
        df = parse_dates(df, date_columns)

        # Save all data to pickle
        save_all_data_to_pickle(df, pkl_dir)

        # Data transformation and model evaluation
        print("Transforming Data")
        df_imputed, label_encoders, transformation_report = transform_data(df)

        # Split data and run models
        # ... (rest of your existing code for model evaluation) ...

        # Generate and save plots
        generate_and_save_plots(df, plots_dir)

        # Create analysis document
        create_analysis_document(
            analysis_dir, df, results, plots_dir, transformation_report
        )

    except Exception as e:
        logging.error(f"Main analysis failed with error: {e}")


def load_file(file_path):
    if file_path.endswith(".csv"):
        return pd.read_csv(file_path)
    elif file_path.endswith(".xlsx"):
        return pd.read_excel(file_path)
    else:
        print("Unsupported file format. Please provide a CSV or Excel file.")
        return pd.DataFrame()


if __name__ == "__main__":
    run_analysis(
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\Datasets\\laptop_prices.csv",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts1",
    )
