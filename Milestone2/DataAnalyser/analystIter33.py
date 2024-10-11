import os
import pickle

import numpy as np
import pandas as pd
import plotly.express as px
from docx import Document
from sklearn.decomposition import PCA
from sklearn.ensemble import (
    RandomForestClassifier,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.svm import SVC


# Function to create directories
def create_directories(output_directory, dataset_name):
    """Create directories for analysis reports and plots."""
    analysis_dir = os.path.join(output_directory, dataset_name, "analysis")
    plots_dir = os.path.join(output_directory, dataset_name, "plots")
    os.makedirs(analysis_dir, exist_ok=True)
    os.makedirs(plots_dir, exist_ok=True)
    return analysis_dir, plots_dir


# Function to save plots as HTML
def save_plot_as_html(plot, plot_name, plots_dir):
    """Save plot as an interactive HTML file."""
    plot_path = os.path.join(plots_dir, f"{plot_name}.html")
    plot.write_html(plot_path)


# Function to read the file (CSV or Excel)
def load_file(file_path):
    """Loads a CSV or Excel file into a DataFrame."""
    try:
        if file_path.endswith(".csv"):
            df = pd.read_csv(file_path)
        elif file_path.endswith(".xlsx"):
            df = pd.read_excel(file_path)
        else:
            raise ValueError(
                "Unsupported file format. Please provide a CSV or Excel file."
            )
        # Remove spaces in column names
        df.columns = df.columns.str.replace(" ", "_")
    except Exception as e:
        print(f"Error loading file {file_path}: {e}")
        df = pd.DataFrame()  # Return empty DataFrame if file load fails
    return df


# Function to suggest the target column using options
def suggest_target_column(df, learning=True):
    """Suggest the target column using simple heuristics and NLP, or learn from previous analyses."""
    target_column = None

    if learning and os.path.exists("target_column_learning.pickle"):
        # Load pickle file with learned target column suggestions
        with open("target_column_learning.pickle", "rb") as f:
            learned_data = pickle.load(f)
        for column, count in learned_data.items():
            if column in df.columns:
                print(f"Learned suggestion: {column} (appeared {count} times)")
        return learned_data

    # First: Check for columns that contain words like "price", "target", or "label"
    for column in df.columns:
        column_name = column.lower()
        if "price" in column_name or "target" in column_name or "label" in column_name:
            target_column = column
            break

    if not target_column:
        print("Please select the target column from the following list:")
        for idx, column in enumerate(df.columns):
            print(f"{idx}: {column}")
        idx = int(input("Enter the index of the target column: "))
        target_column = df.columns[idx]

    return target_column


# Function to perform data transformation
def transform_data(df, is_numerical=True):
    """Transform the dataset for model usage based on the type (numerical or categorical)."""
    transformation_report = []

    # Handling missing values using SimpleImputer (mean strategy for numerical columns)
    imputer = SimpleImputer(strategy="mean")
    df_imputed = df.copy()  # To avoid SettingWithCopyWarning
    numerical_columns = df.select_dtypes(include=[np.number]).columns

    if is_numerical:
        # Apply transformations for numerical data
        df_imputed[numerical_columns] = imputer.fit_transform(
            df_imputed[numerical_columns]
        )
        transformation_report.append(
            "Missing values were imputed using the SimpleImputer with mean strategy."
        )

        # Check if normalization is necessary, avoid if it leads to 0 values
        normalizer = Normalizer()
        if (
            df_imputed[numerical_columns].max().max() != 0
        ):  # Check if normalization is needed
            df_imputed[numerical_columns] = normalizer.fit_transform(
                df_imputed[numerical_columns]
            )
            transformation_report.append(
                "Normalization was applied to numerical features."
            )
        else:
            transformation_report.append(
                "Normalization skipped due to zero values in data."
            )
    else:
        # Apply transformations for categorical data
        label_encoders = {}
        categorical_columns = df.select_dtypes(include=["object"]).columns
        for column in categorical_columns:
            le = LabelEncoder()
            df_imputed[column] = le.fit_transform(df[column].astype(str))
            label_encoders[column] = le
        transformation_report.append(
            "Label Encoding was applied to categorical variables."
        )

    return df_imputed, transformation_report


# Apply PCA to the transformed data
def apply_pca(X, n_components=0.95):
    """Apply PCA for dimensionality reduction."""
    try:
        pca = PCA(n_components=n_components)
        X_pca = pca.fit_transform(X)
        return X_pca, pca
    except Exception as e:
        print(f"Error during PCA: {e}")
        return X, None


# Generate useful plots and save as HTML
def generate_plots(df, target_column, plots_dir):
    """Generate useful plots (feature distributions, PCA, correlation heatmap)."""
    # Feature distribution plot (for numerical columns)
    for column in df.select_dtypes(include=[np.number]).columns:
        fig = px.histogram(df, x=column, title=f"Distribution of {column}")
        save_plot_as_html(fig, f"feature_distribution_{column}", plots_dir)

    # Correlation heatmap
    corr_matrix = df.corr()
    fig = px.imshow(corr_matrix, title="Correlation Heatmap")
    save_plot_as_html(fig, "correlation_heatmap", plots_dir)

    # PCA plot (for dimensionality reduction)
    X = df.drop(columns=[target_column])
    X_pca, _ = apply_pca(X, n_components=2)
    pca_df = pd.DataFrame(X_pca, columns=["PCA1", "PCA2"])
    pca_df[target_column] = df[target_column]
    fig = px.scatter(pca_df, x="PCA1", y="PCA2", color=target_column, title="PCA Plot")
    save_plot_as_html(fig, "pca_plot", plots_dir)

    # Check for datetime columns and generate time series plots if available
    datetime_columns = df.select_dtypes(include=["datetime"]).columns
    if len(datetime_columns) > 0:
        for column in datetime_columns:
            df[column] = pd.to_datetime(
                df[column], errors="coerce"
            )  # Convert to datetime
            df_sorted = df.sort_values(by=column)
            fig = px.line(
                df_sorted,
                x=column,
                y=target_column,
                title=f"Time Series Plot for {target_column}",
            )
            save_plot_as_html(fig, f"time_series_{target_column}", plots_dir)


# Save transformation methods to Word document
def save_transformation_methods(df, doc_filename, transformation_report):
    """Saves the transformation methods used to a Word document."""
    doc = Document()
    doc.add_heading("Data Transformation Report", 0)

    doc.add_heading("1. Missing Data Imputation:", level=1)
    doc.add_paragraph(transformation_report[0])

    doc.add_heading("2. Label Encoding for Categorical Columns:", level=1)
    doc.add_paragraph(transformation_report[2])

    doc.add_heading("3. Feature Normalization:", level=1)
    doc.add_paragraph(transformation_report[1])

    doc.add_heading("Transformed Data Preview:", level=1)
    doc.add_paragraph(str(df.head()))

    doc.save(doc_filename)


# Function to log training errors
def log_training_errors(message):
    """Log training errors to a file."""
    with open("training_log.txt", "a") as log_file:
        log_file.write(message + "\n")


# Main function to invoke all processes
def run_analysis(file_path, output_directory):
    """Runs the entire pipeline of loading, transforming, analyzing, and saving reports."""
    try:
        df = load_file(file_path)
        if df.empty:
            print("No data to process.")
            return

        # Suggest the target column and transform the data
        target_column = suggest_target_column(df)
        print(f"target_column : {target_column}")

        print("Transforming Data")

        # Apply transformations for numerical models on data1
        data1_imputed, num_transformation_report = transform_data(df, is_numerical=True)

        # Create directories
        dataset_name = os.path.basename(file_path).split(".")[0]
        analysis_dir, plots_dir = create_directories(output_directory, dataset_name)

        # Split the data into training and testing sets
        X1 = data1_imputed.drop(columns=[target_column])
        y1 = data1_imputed[target_column]
        X_train1, X_test1, y_train1, y_test1 = train_test_split(
            X1, y1, test_size=0.3, random_state=42
        )

        # Try auto-training with multiple models and pick the best one
        models = {
            "Logistic Regression": LogisticRegression(max_iter=200),
            "Random Forest": RandomForestClassifier(),
            "SVM": SVC(),
        }

        best_model = None
        best_score = 0

        for model_name, model in models.items():
            try:
                model.fit(X_train1, y_train1)
                score = model.score(X_test1, y_test1)
                if score > best_score:
                    best_score = score
                    best_model = model
            except Exception as e:
                log_training_errors(f"Error in training {model_name}: {e}")

        print(f"Best model: {best_model.__class__.__name__} with score: {best_score}")

        # Save results and generate analysis
        generate_plots(df, target_column, plots_dir)
        save_transformation_methods(
            df,
            os.path.join(analysis_dir, "transformation_report.docx"),
            num_transformation_report,
        )

        print("Analysis completed and saved.")

    except Exception as e:
        print(f"Error during analysis: {e}")
        log_training_errors(f"Error during analysis: {e}")


if __name__ == "__main__":
    run_analysis(
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\Datasets\\StudentPerformanceFactors.csv",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts",
    )
