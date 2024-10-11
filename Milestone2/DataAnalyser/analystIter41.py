import logging
import os
import pickle
import re

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
from docx import Document
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    RandomForestRegressor,
    GradientBoostingRegressor,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.metrics import (
    confusion_matrix,
    classification_report,
    mean_absolute_error,
    mean_squared_error,
    r2_score,
)
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.svm import SVC, SVR
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor

# Setup Plotly renderer
pio.renderers.default = "browser"

# Global analysis log
analysis_log = {
    "steps": [],
    "reasons": [],
}


def log_step(step, reason):
    analysis_log["steps"].append(step)
    analysis_log["reasons"].append(reason)


# Setup logging
def setup_logging(log_file):
    logging.basicConfig(
        filename=log_file,
        level=logging.ERROR,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


# Check for symbols and spaces in column names
def check_column_names(df):
    """Check for symbols and spaces in DataFrame column names."""
    invalid_columns = [col for col in df.columns if re.search(r"[^a-zA-Z0-9_]", col)]
    if invalid_columns:
        print(f"Warning: Columns with symbols/spaces detected: {invalid_columns}")
        df.columns = df.columns.str.replace(r"[^a-zA-Z0-9_]", "_", regex=True)
    return df


# Create required output directories
def create_directories(output_directory, dataset_name):
    dirs = ["analysis", "plots", "time_series"]
    created_dirs = {}
    for d in dirs:
        created_dirs[d] = os.path.join(output_directory, dataset_name, d)
        os.makedirs(created_dirs[d], exist_ok=True)
    return created_dirs


# Load previous data from pickle
def load_pickle_data(pkl_dir):
    pkl_file = os.path.join(pkl_dir, "data.pkl")
    if os.path.exists(pkl_file):
        with open(pkl_file, "rb") as f:
            try:
                return pickle.load(f)
            except EOFError:
                return None
    return None


# Save model, vectorizer, transformation steps, and dataset metadata to pickle
def save_to_pickle(
    df,
    pkl_dir,
    target_column,
    best_model=None,
    vectorizer=None,
    transformation_steps=None,
):
    dataset_metadata = {
        "target_column": target_column,
        "shape": df.shape,
        "columns": df.columns.tolist(),
        "best_model": best_model,
        "vectorizer": vectorizer,
        "transformation_steps": transformation_steps,
    }
    pkl_file = os.path.join(pkl_dir, "data.pkl")
    with open(pkl_file, "ab") as f:
        pickle.dump(dataset_metadata, f)


# Suggest target column
def suggest_target_column(df, pkl_dir):
    log_step(
        "Suggesting target column based on existing columns",
        "To identify the appropriate target for modeling.",
    )
    previous_data = load_pickle_data(pkl_dir)

    if previous_data and "target_column" in previous_data:
        print(f"Previous target column found: {previous_data['target_column']}")
        return previous_data["target_column"]

    for column in df.columns:
        if any(x in column.lower() for x in ["price", "target", "label"]):
            target_column = column
            break
    else:
        target_column = input(f"Choose from: {list(df.columns)}")

    print(f"Suggested target column: {target_column}")
    confirmation = input("Is this okay? (y/n): ")
    if confirmation.lower() != "y":
        target_column = input(
            f"Please choose the target column from: {list(df.columns)}"
        )

    return target_column


# Data transformation
def transform_data(df):
    log_step(
        "Transforming data by imputing and normalizing",
        "To handle missing values and ensure data consistency.",
    )
    print("Starting data transformation...")
    imputer = SimpleImputer(strategy="mean")
    df_imputed = df.copy()

    numerical_columns = df.select_dtypes(include=[np.number]).columns
    df_imputed[numerical_columns] = imputer.fit_transform(df_imputed[numerical_columns])

    # Normalization
    if df_imputed[numerical_columns].max().max() != 0:
        normalizer = Normalizer()
        df_imputed[numerical_columns] = normalizer.fit_transform(
            df_imputed[numerical_columns]
        )

    # Label encoding for categorical variables
    label_encoders = {}
    categorical_columns = df.select_dtypes(include=["object"]).columns
    for column in categorical_columns:
        le = LabelEncoder()
        df_imputed[column] = le.fit_transform(df_imputed[column].astype(str))
        label_encoders[column] = le

    print("Data transformation completed.")
    return df_imputed, label_encoders


# Evaluate models
def evaluate_models(df, target_column, plots_dir):
    log_step(
        "Evaluating models based on the target column type",
        "To assess performance using suitable algorithms.",
    )
    results = {}
    X = df.drop(columns=[target_column])
    y = df[target_column]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )

    models = select_models(y)

    for model_name, model in models.items():
        try:
            model.fit(X_train, y_train)
            predictions = model.predict(X_test)
            results[model_name] = evaluate_model(
                y, y_test, predictions, model_name, plots_dir
            )
        except Exception as e:
            log_error(model_name, e)

    return results


def select_models(y):
    if y.dtype in [np.int64, np.float64] and len(y.unique()) > 2:
        return {  # Regression
            "Linear Regression": LinearRegression(),
            "Random Forest Regressor": RandomForestRegressor(),
            "Gradient Boosting Regressor": GradientBoostingRegressor(),
            "K-Neighbors Regressor": KNeighborsRegressor(),
            "Support Vector Regressor": SVR(),
            "Decision Tree Regressor": DecisionTreeRegressor(),
        }
    else:  # Classification
        return {
            "Logistic Regression": LogisticRegression(max_iter=1000),
            "Random Forest": RandomForestClassifier(),
            "Gradient Boosting": GradientBoostingClassifier(),
            "K-Neighbors": KNeighborsClassifier(),
            "Support Vector Classifier": SVC(),
            "Naive Bayes": GaussianNB(),
            "Decision Tree": DecisionTreeClassifier(),
        }


def evaluate_model(y, y_test, predictions, model_name, plots_dir):
    if y.dtype in [np.int64, np.float64] and len(y.unique()) > 2:  # Regression
        return regression_metrics(y_test, predictions, model_name)
    else:  # Classification
        return classification_metrics(y_test, predictions, model_name, plots_dir)


def regression_metrics(y_test, predictions, model_name):
    mae = mean_absolute_error(y_test, predictions)
    mse = mean_squared_error(y_test, predictions)
    r2 = r2_score(y_test, predictions)

    print(f"{model_name} Metrics:\nMAE: {mae}, MSE: {mse}, R²: {r2}")
    return {"MAE": mae, "MSE": mse, "R²": r2}


def classification_metrics(y_test, predictions, model_name, plots_dir):
    conf_matrix = confusion_matrix(y_test, predictions)
    class_report = classification_report(y_test, predictions)

    print(f"Confusion Matrix for {model_name}:\n{conf_matrix}")
    print(f"Classification Report for {model_name}:\n{class_report}")

    # Plotting the confusion matrix with Plotly
    fig = px.imshow(
        conf_matrix,
        text_auto=True,
        title=f"Confusion Matrix: {model_name}",
        labels=dict(x="Predicted Label", y="True Label"),
        color_continuous_scale="Blues",
    )
    fig.write_html(os.path.join(plots_dir, f"confusion_matrix_{model_name}.html"))

    return {"Confusion Matrix": conf_matrix, "Classification Report": class_report}


# Generate and save plots
def generate_and_save_plots(df, plots_dir, target_column):
    log_step(
        "Generating and saving plots for data analysis",
        "To visualize data distributions and relationships.",
    )
    print("Generating plots...")
    generate_histograms(df, plots_dir)
    generate_correlation_matrix(df, plots_dir)
    generate_target_feature_plots(df, target_column, plots_dir)


def generate_histograms(df, plots_dir):
    numerical_cols = df.select_dtypes(include=[np.number]).columns
    for col in numerical_cols:
        fig = px.histogram(
            df,
            x=col,
            title=f"Histogram of {col}",
            labels={col: col},
        )
        fig.write_html(os.path.join(plots_dir, f"histogram_{col}.html"))


def generate_correlation_matrix(df, plots_dir):
    fig = px.imshow(
        df.corr(),
        text_auto=True,
        title="Correlation Matrix",
        labels=dict(x="Features", y="Features"),
        color_continuous_scale="Viridis",
    )
    fig.write_html(os.path.join(plots_dir, "correlation_matrix.html"))


def generate_target_feature_plots(df, target_column, plots_dir):
    for col in df.columns:
        if col != target_column:
            fig = px.scatter(
                df,
                x=col,
                y=target_column,
                title=f"Scatter Plot: {col} vs {target_column}",
                labels={col: col, target_column: target_column},
            )
            fig.write_html(
                os.path.join(plots_dir, f"scatter_{col}_vs_{target_column}.html")
            )


# Save the analysis log and results to a Word document
def save_analysis_log(log_data, output_dir, filename="analysis_log.docx"):
    doc = Document()
    doc.add_heading("Data Analysis Log", 0)

    for i, (step, reason) in enumerate(zip(log_data["steps"], log_data["reasons"])):
        doc.add_paragraph(f"{i + 1}. {step}: {reason}")

    doc.save(os.path.join(output_dir, filename))


# Error logging
def log_error(model_name, error):
    logging.error(f"Error in {model_name}: {error}")
    print(f"Error in {model_name}: {error}. Continuing with other models.")


# Core pipeline for data analysis
def data_analysis_pipeline(dataset_path, output_directory, log_file, pkl_dir):
    setup_logging(log_file)

    df = pd.read_csv(dataset_path)
    df = check_column_names(df)

    dataset_name = os.path.basename(dataset_path).split(".")[0]
    created_dirs = create_directories(output_directory, dataset_name)

    target_column = suggest_target_column(df, pkl_dir)

    df_transformed, label_encoders = transform_data(df)

    model_results = evaluate_models(
        df_transformed, target_column, created_dirs["plots"]
    )

    generate_and_save_plots(df_transformed, created_dirs["plots"], target_column)

    best_model = max(model_results, key=lambda x: model_results[x]["MAE"])
    save_to_pickle(df_transformed, pkl_dir, target_column, best_model=best_model)

    save_analysis_log(analysis_log, created_dirs["analysis"])

    print(
        f"Analysis completed for {dataset_name}. Results saved to {output_directory}."
    )


if __name__ == "__main__":
    data_analysis_pipeline(
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\Datasets\\laptop_prices.csv",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataReadOuts1",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\DataAnalyser\\training_log.txt",
        "C:\\Users\\HimakarRaju\\Desktop\\Milestone2\\pkls",
    )
