import os
import shutil
import pickle
import pandas as pd
import numpy as np
import logging
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import (
    StandardScaler,
    LabelEncoder,
    OneHotEncoder,
    MinMaxScaler,
    PolynomialFeatures,
)
from sklearn.linear_model import LinearRegression, LogisticRegression, Ridge, Lasso
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    AdaBoostClassifier,
    VotingClassifier,
)
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import GaussianNB
from sklearn.cluster import KMeans
from sklearn.metrics import (
    accuracy_score,
    mean_squared_error,
    r2_score,
    confusion_matrix,
)
from plotly import express as px
from plotly import graph_objects as go
from docx import Document
from docx.shared import Inches
from sklearn.feature_extraction.text import CountVectorizer

# Ensure the errors_log directory exists
log_dir = "root/errors_log"
os.makedirs(log_dir, exist_ok=True)

# Setting up logging
logging.basicConfig(
    filename=os.path.join(log_dir, "errors.log"),
    level=logging.ERROR,
    format="%(asctime)s:%(levelname)s:%(message)s",
)


class Analyser:
    def __init__(self):
        self.dataset_path = None
        self.files = []
        self.data = None
        self.document = Document()
        self.model_results = []
        self.best_model = None
        self.output_folder = None
        self.target_column = None
        self.analysis_log = {"steps": [], "reasons": []}

    def add_requirement(self):
        try:
            print("Step 1: Adding requirements.")
            self.dataset_path = input(
                "Enter the path of the dataset (CSV/XLSX file or folder containing them): "
            ).strip()
            self.output_folder = input("Enter the path of the output folder: ").strip()
            os.makedirs(self.output_folder, exist_ok=True)
            print(f"Dataset path set to: {self.dataset_path}")
            print("Requirements added successfully.")
        except Exception as e:
            logging.error(f"Error in add_requirement: {e}")

    def create_directories(self, dataset_name):
        """Create necessary directories for saving outputs."""
        base_dir = os.path.join(self.output_folder, f"{dataset_name}_analysis")
        os.makedirs(base_dir, exist_ok=True)
        dirs = ["Analysis", "Plots", "SVGs"]
        created_dirs = {d: os.path.join(base_dir, d) for d in dirs}
        for dir_path in created_dirs.values():
            os.makedirs(dir_path, exist_ok=True)
        return created_dirs

    def run_analysis(self):
        try:
            self.collect_files(self.dataset_path)
            if not self.files:
                print("No valid files found.")
            else:
                for file in self.files:
                    print(f"Analyzing file: {file}")
                    self.document.add_heading(f"Analysis for {file}", level=1)
                    self.clean_data(file)
                    self.select_target_column()
                    self.vectorize_text()
                    self.apply_models()  # This method needs to append to self.model_results
                    self.generate_and_save_plots(file)

                    # Call finalize_analysis after processing each file
                    self.finalize_analysis(file)

                    self.save_documentation(file)
                    self.save_pickles(file)
        except Exception as e:
            logging.error(f"Error in run_analysis: {e}")

    def collect_files(self, path):
        try:
            if os.path.isfile(path) and path.endswith((".csv", ".xlsx")):
                self.files.append(path)
            elif os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for file in files:
                        if file.endswith((".csv", ".xlsx")):
                            self.files.append(os.path.join(root, file))
            else:
                print(f"Invalid path: {path}")
        except Exception as e:
            logging.error(f"Error in collect_files: {e}")

    def clean_data(self, file_path):
        try:
            print(f"Step 2: Cleaning data from {file_path}.")
            if file_path.endswith(".csv"):
                self.data = pd.read_csv(file_path)
            elif file_path.endswith(".xlsx"):
                self.data = pd.read_excel(file_path)

            self.numerical_data = self.data.select_dtypes(include=["number"])
            self.categorical_data = self.data.select_dtypes(exclude=["number"])

            print("Data cleaned successfully.")
            self.document.add_paragraph("Data Cleaning Steps:")
            self.document.add_paragraph(
                f"Numerical Data:\n{self.numerical_data.head().to_string()}"
            )
            self.document.add_paragraph(
                f"Categorical Data:\n{self.categorical_data.head().to_string()}"
            )
        except Exception as e:
            logging.error(f"Error in clean_data: {e}")

    def select_target_column(self):
        try:
            print("Step 3: Selecting target column.")
            self.target_column = None
            for col in self.data.columns:
                if "target" in col.lower() or "label" in col.lower():
                    self.target_column = col
                    break

            if not self.target_column:
                print("No target column found automatically. Available columns:")
                print(self.data.columns.tolist())
                self.target_column = input("Please specify the target column: ").strip()

            print(f"Target column set to: {self.target_column}")
            self.document.add_paragraph(f"Selected target column: {self.target_column}")
            print("Target column selection completed.")
        except Exception as e:
            logging.error(f"Error in select_target_column: {e}")

    def apply_models(self):
        try:
            print("Step 4: Applying models.")
            X = self.data.drop(columns=self.target_column)
            y = self.data[self.target_column]

            models = [
                {"name": "Linear Regression", "model": LinearRegression()},
                {"name": "Ridge Regression", "model": Ridge()},
                {"name": "Lasso Regression", "model": Lasso()},
                {
                    "name": "Logistic Regression",
                    "model": LogisticRegression(max_iter=200),
                },
                {"name": "Decision Tree", "model": DecisionTreeClassifier()},
                {"name": "Random Forest", "model": RandomForestClassifier()},
                {"name": "Gradient Boosting", "model": GradientBoostingClassifier()},
                {"name": "AdaBoost", "model": AdaBoostClassifier()},
                {"name": "SVC", "model": SVC()},
                {"name": "K Neighbors", "model": KNeighborsClassifier()},
                {"name": "GaussianNB", "model": GaussianNB()},
                {"name": "KMeans", "model": KMeans(n_clusters=3)},
                {
                    "name": "Voting Classifier",
                    "model": VotingClassifier(
                        estimators=[
                            ("lr", LogisticRegression(max_iter=200)),
                            ("rf", RandomForestClassifier()),
                            ("svc", SVC()),
                        ]
                    ),
                },
            ]

            results_list = []
            best_score = float("-inf")

            for model_info in models:
                model = model_info["model"]
                try:
                    print(f"Training model: {model_info['name']}.")
                    X_transformed, y_transformed = self.transform_data(model, X, y)
                    X_train, X_test, y_train, y_test = train_test_split(
                        X_transformed, y_transformed, test_size=0.2, random_state=42
                    )

                    # Fit model and evaluate
                    model.fit(X_train, y_train)
                    train_score = model.score(X_train, y_train)
                    test_score = model.score(X_test, y_test)
                    y_pred = model.predict(X_test)

                    # Calculate metrics
                    r2 = r2_score(y_test, y_pred)
                    mae = np.mean(np.abs(y_test - y_pred))
                    mse = mean_squared_error(y_test, y_pred)

                    # Store results
                    result = {
                        "name": model_info["name"],
                        "train": train_score,
                        "test": test_score,
                        "r2score": r2,
                        "absolute": mae,
                        "squared": mse,
                    }
                    results_list.append(result)
                    self.model_results.append((model, result))

                    if test_score > best_score:
                        best_score = test_score
                        self.best_model = model

                    print(f"Model {model_info['name']} trained successfully.")

                except Exception as model_error:
                    logging.error(
                        f"Error while applying model {model_info['name']}: {model_error}"
                    )

            print("Model evaluation completed.")
            # Save results to DataFrame and document
            results_df = pd.DataFrame(results_list)
            results_df.set_index("name", inplace=True)
            results_df.sort_values(by="test", ascending=False, inplace=True)

            self.document.add_heading("Model Evaluation Results", level=2)
            for index, row in results_df.iterrows():
                self.document.add_paragraph(
                    f"{index}: Train Score: {row['train']:.4f}, Test Score: {row['test']:.4f}, "
                    f"R² Score: {row['r2score']:.4f}, Mean Absolute Error: {row['absolute']:.4f}, "
                    f"Mean Squared Error: {row['squared']:.4f}"
                )

        except Exception as e:
            logging.error(f"Error in apply_models: {e}")

    def generate_and_save_plots(self, file_name):
        """Generate and save plots for data analysis."""
        print("Generating plots...")
        created_dirs = self.create_directories(
            os.path.basename(file_name).split(".")[0]
        )
        plots_dir = created_dirs["Plots"]
        svg_dir = created_dirs["SVGs"]
        self.generate_target_feature_plots(plots_dir, svg_dir)

    def generate_target_feature_plots(self, plots_dir, svg_dir):
        """Generate plots for the target variable against all feature variables."""
        for col in self.data.drop(columns=[self.target_column]).columns:
            fig = go.Figure()

            # Scatter plot
            fig.add_trace(
                go.Scatter(
                    x=self.data[col],
                    y=self.data[self.target_column],
                    mode="markers",
                    name="Data Points",
                )
            )

            # Fit trend lines
            # Linear regression
            linear_model = LinearRegression()
            linear_model.fit(self.data[[col]], self.data[self.target_column])
            linear_pred = linear_model.predict(self.data[[col]])
            fig.add_trace(
                go.Scatter(
                    x=self.data[col], y=linear_pred, mode="lines", name="Linear Trend"
                )
            )

            # Polynomial regression
            poly_features = PolynomialFeatures(degree=2)
            X_poly = poly_features.fit_transform(self.data[[col]])
            poly_model = LinearRegression()
            poly_model.fit(X_poly, self.data[self.target_column])
            poly_pred = poly_model.predict(X_poly)
            fig.add_trace(
                go.Scatter(
                    x=self.data[col], y=poly_pred, mode="lines", name="Polynomial Trend"
                )
            )

            # Determine the best fit based on R²
            linear_r2 = r2_score(self.data[self.target_column], linear_pred)
            poly_r2 = r2_score(self.data[self.target_column], poly_pred)

            if linear_r2 > poly_r2:
                chosen_model = "Linear"
                chosen_pred = linear_pred
                chosen_r2 = linear_r2
            else:
                chosen_model = "Polynomial"
                chosen_pred = poly_pred
                chosen_r2 = poly_r2

            # Add the chosen trend line
            fig.add_trace(
                go.Scatter(
                    x=self.data[col],
                    y=chosen_pred,
                    mode="lines",
                    name=f"Chosen {chosen_model} Trend",
                )
            )

            # Add title and labels
            fig.update_layout(
                title=f"{self.target_column} vs {col} (Chosen: {chosen_model} Trend)",
                xaxis_title=col,
                yaxis_title=self.target_column,
            )

            # Save plots
            fig.write_html(
                os.path.join(plots_dir, f"{self.target_column}_vs_{col}.html")
            )
            fig.write_image(
                os.path.join(plots_dir, f"{self.target_column}_vs_{col}.png")
            )
            fig.write_image(os.path.join(svg_dir, f"{self.target_column}_vs_{col}.svg"))

            # Log the chosen trend line
            self.analysis_log["steps"].append(
                f"Trend Line for {self.target_column} vs {col}"
            )
            self.analysis_log["reasons"].append(
                f"Chosen {chosen_model} trend due to R² value: {chosen_r2:.4f}"
            )

    def finalize_analysis(self, file_name):
        results = {
            model.__class__.__name__: metrics for model, metrics in self.model_results
        }
        plots_dir = os.path.join(
            self.output_folder,
            f"{os.path.basename(file_name).split('.')[0]}_analysis",
            "b) plots",
        )
        analysis_dir = os.path.join(
            self.output_folder,
            f"{os.path.basename(file_name).split('.')[0]}_analysis",
            "Analysis",
        )

        # Log analysis steps
        print("Analysis Steps and Reasons:")
        for step, reason in zip(
            self.analysis_log["steps"], self.analysis_log["reasons"]
        ):
            print(f"Step: {step}\nReason: {reason}\n")

        # Create Word document
        doc = Document()
        doc.add_heading("Analysis Report", level=1)

        # Methods section
        doc.add_heading("Methods Used", level=2)
        for step, reason in zip(
            self.analysis_log["steps"], self.analysis_log["reasons"]
        ):
            doc.add_paragraph(f"{step}: {reason}")

        # Results section
        doc.add_heading("Model Results", level=2)
        for model_name, metrics in results.items():
            doc.add_heading(model_name, level=3)
            if "score" in metrics:
                doc.add_paragraph(f"Score: {metrics['score']:.4f}")
            else:
                doc.add_paragraph(f"MSE: {metrics['mse']:.4f}")
                doc.add_paragraph(f"R²: {metrics['r2']:.4f}")

        # Plot Understanding Section
        doc.add_heading("Understanding of Plots", level=2)
        for col in self.data.drop(columns=[self.target_column]).columns:
            doc.add_heading(f"{self.target_column} vs {col}", level=3)
            doc.add_paragraph(
                f"Chosen trend line: {self.analysis_log['steps'][self.data.columns.get_loc(col)]}"
            )
            doc.add_paragraph(
                self.analysis_log["reasons"][self.data.columns.get_loc(col)]
            )
            img_path = os.path.join(plots_dir, f"{self.target_column}_vs_{col}.png")
            doc.add_picture(img_path, width=Inches(5))
            doc.add_paragraph(
                f"Analysis of the plot: The relationship between {col} and {self.target_column} suggests..."
            )  # Placeholder text for analysis

        # Best model section
        best_model_info = self.determine_best_model(results)
        doc.add_heading("Best Model", level=2)
        doc.add_paragraph(
            f"The best model is: {best_model_info['name']} with score: {best_model_info['score']:.4f}"
        )

        # Save the document in the analysis folder
        doc_file_path = os.path.join(analysis_dir, "Analysis_Report.docx")
        doc.save(doc_file_path)

        print(f"Report saved to {doc_file_path}")

    def determine_best_model(self, results):
        best_model_name = None
        best_score = float("-inf")
        for model_name, metrics in results.items():
            if "score" in metrics and metrics["score"] > best_score:
                best_score = metrics["score"]
                best_model_name = model_name
        return {"name": best_model_name, "score": best_score}

    def save_documentation(self, file_name):
        try:
            doc_path = os.path.join(
                self.output_folder,
                f"{os.path.basename(file_name).split('.')[0]}_analysis",
            )
            os.makedirs(doc_path, exist_ok=True)
            self.document.save(os.path.join(doc_path, "analysis_report.docx"))
            print(
                f"Documentation saved at {os.path.join(doc_path, 'analysis_report.docx')}"
            )
        except Exception as e:
            logging.error(f"Error in save_documentation: {e}")

    def save_pickles(self, file_name):
        try:
            pkl_path = os.path.join(
                self.output_folder,
                f"{os.path.basename(file_name).split('.')[0]}_analysis",
            )
            os.makedirs(pkl_path, exist_ok=True)
            for model, _ in self.model_results:
                with open(
                    os.path.join(pkl_path, f"{model.__class__.__name__}.pkl"), "wb"
                ) as f:
                    pickle.dump(model, f)
            print(f"Pickles saved at {pkl_path}")
        except Exception as e:
            logging.error(f"Error in save_pickles: {e}")


# Usage
if __name__ == "__main__":
    try:
        analyser = Analyser()
        analyser.add_requirement()
        analyser.run_analysis()
    except Exception as e:
        logging.error(f"Critical error in main execution: {e}")
