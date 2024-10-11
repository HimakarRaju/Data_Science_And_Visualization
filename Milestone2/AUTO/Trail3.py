import os
import pickle
import pandas as pd
import numpy as np
import logging
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import (
    StandardScaler,
    LabelEncoder,
    OneHotEncoder,
    MinMaxScaler,
    Binarizer,
    PolynomialFeatures,
)
from sklearn.linear_model import LinearRegression, LogisticRegression, Ridge, Lasso
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    AdaBoostClassifier,
    VotingClassifier,
)
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import GaussianNB
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.metrics import accuracy_score, mean_squared_error, r2_score
from plotly import express as px
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Ensure the errors_log directory exists
log_dir = "root/errors_log"
os.makedirs(log_dir, exist_ok=True)

# Setting up logging
logging.basicConfig(
    filename=os.path.join(log_dir, "errors.log"),
    level=logging.ERROR,
    format="%(asctime)s:%(levelname)s:%(message)s",
)


class Analyser:
    def __init__(self):
        self.dataset_path = None
        self.files = []
        self.data = None
        self.document = Document()
        self.model_results = []
        self.best_model = None
        self.output_folder = None
        self.target_column = None

    def add_requirement(self):
        try:
            self.dataset_path = input(
                "Enter the path of the dataset (CSV/XLSX file or folder containing them): "
            ).strip()
            self.output_folder = input("Enter the path of the output folder: ").strip()
            os.makedirs(self.output_folder, exist_ok=True)
            print(f"Dataset path set to: {self.dataset_path}")
        except Exception as e:
            logging.error(f"Error in add_requirement: {e}")

    def run_analysis(self):
        try:
            self.collect_files(self.dataset_path)
            if not self.files:
                print("No valid files found.")
            else:
                for file in self.files:
                    print(f"Analyzing file: {file}")
                    self.document.add_heading(f"Analysis for {file}", level=1)
                    self.clean_data(file)
                    self.select_target_column()
                    self.vectorize_text()
                    self.apply_models()
                    self.plot_graphs(file)
                    self.save_documentation(file)
                    self.save_pickles(file)
        except Exception as e:
            logging.error(f"Error in run_analysis: {e}")

    def collect_files(self, path):
        try:
            if os.path.isfile(path) and path.endswith((".csv", ".xlsx")):
                self.files.append(path)
            elif os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for file in files:
                        if file.endswith((".csv", ".xlsx")):
                            self.files.append(os.path.join(root, file))
            else:
                print(f"Invalid path: {path}")
        except Exception as e:
            logging.error(f"Error in collect_files: {e}")

    def clean_data(self, file_path):
        try:
            if file_path.endswith(".csv"):
                self.data = pd.read_csv(file_path)
            elif file_path.endswith(".xlsx"):
                self.data = pd.read_excel(file_path)

            self.numerical_data = self.data.select_dtypes(include=["number"])
            self.categorical_data = self.data.select_dtypes(exclude=["number"])

            self.document.add_paragraph("Data Cleaning Steps:")
            self.document.add_paragraph(
                f"Numerical Data:\n{self.numerical_data.head().to_string()}"
            )
            self.document.add_paragraph(
                f"Categorical Data:\n{self.categorical_data.head().to_string()}"
            )
        except Exception as e:
            logging.error(f"Error in clean_data: {e}")

    def select_target_column(self):
        try:
            # Attempt to find a suitable target column automatically
            self.target_column = None
            for col in self.data.columns:
                if "target" in col.lower() or "label" in col.lower():
                    self.target_column = col
                    break

            if not self.target_column:
                # Fallback: ask user to select if none found
                print("No target column found automatically. Available columns:")
                print(self.data.columns.tolist())
                self.target_column = input("Please specify the target column: ").strip()

            print(f"Target column set to: {self.target_column}")
            self.document.add_paragraph(f"Selected target column: {self.target_column}")
        except Exception as e:
            logging.error(f"Error in select_target_column: {e}")

    def vectorize_text(self):
        try:
            if self.categorical_data.shape[1] > 0:
                vectorizer = CountVectorizer()
                text_matrix = vectorizer.fit_transform(
                    self.data[self.categorical_data.columns]
                    .astype(str)
                    .values.flatten()
                )
                self.data = pd.DataFrame(
                    text_matrix.toarray(), columns=vectorizer.get_feature_names_out()
                )
                print("Text data vectorized.")
        except Exception as e:
            logging.error(f"Error in vectorize_text: {e}")

    def transform_data(self, model, X, y=None):
        try:
            if isinstance(
                model,
                (
                    LinearRegression,
                    Ridge,
                    Lasso,
                    LogisticRegression,
                    DecisionTreeClassifier,
                    RandomForestClassifier,
                    GradientBoostingClassifier,
                    AdaBoostClassifier,
                    SVC,
                    KNeighborsClassifier,
                ),
            ):
                scaler = StandardScaler()
                X_transformed = scaler.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, KMeans):
                scaler = MinMaxScaler()
                X_transformed = scaler.fit_transform(X)
                return X_transformed
            elif isinstance(model, GaussianNB):
                label_encoder = LabelEncoder()
                for col in X.select_dtypes(include=["object"]).columns:
                    X[col] = label_encoder.fit_transform(X[col])
                return X, y
            elif isinstance(model, VotingClassifier):
                encoder = OneHotEncoder()
                X_transformed = encoder.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, PolynomialFeatures):
                poly = PolynomialFeatures(degree=2)
                X_transformed = poly.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, PCA):
                pca = PCA(n_components=2)
                X_transformed = pca.fit_transform(X)
                return X_transformed, y
            return X, y
        except Exception as e:
            logging.error(f"Error in transform_data: {e}")
            return X, y  # Return original if transformation fails

    def apply_models(self):
        try:
            X = self.data.drop(columns=self.target_column)
            y = self.data[self.target_column]

            models = [
                LinearRegression(),
                Ridge(),
                Lasso(),
                LogisticRegression(max_iter=200),
                DecisionTreeClassifier(),
                RandomForestClassifier(),
                GradientBoostingClassifier(),
                AdaBoostClassifier(),
                SVC(),
                KNeighborsClassifier(),
                GaussianNB(),
                KMeans(n_clusters=3),
                VotingClassifier(
                    estimators=[
                        ("lr", LogisticRegression(max_iter=200)),
                        ("rf", RandomForestClassifier()),
                        ("svc", SVC()),
                    ]
                ),
            ]

            best_score = float("-inf")

            for model in models:
                try:
                    X_transformed, y_transformed = self.transform_data(model, X, y)
                    X_train, X_test, y_train, y_test = train_test_split(
                        X_transformed, y_transformed, test_size=0.2, random_state=42
                    )
                    model.fit(X_train, y_train)
                    y_pred = model.predict(X_test)

                    if hasattr(model, "score"):
                        score = model.score(X_test, y_test)
                        self.model_results.append((model, {"score": score}))
                        if score > best_score:
                            best_score = score
                            self.best_model = model
                    else:
                        mse = mean_squared_error(y_test, y_pred)
                        r2 = r2_score(y_test, y_pred)
                        self.model_results.append((model, {"mse": mse, "r2": r2}))

                except Exception as model_error:
                    logging.error(
                        f"Error while applying model {model.__class__.__name__}: {model_error}"
                    )

            if self.best_model:
                self.document.add_paragraph(
                    f"Best Model: {self.best_model.__class__.__name__} with score: {best_score}"
                )
                with open("best_model.pkl", "wb") as f:
                    pickle.dump(self.best_model, f)

        except Exception as e:
            logging.error(f"Error in apply_models: {e}")

    def plot_graphs(self, file_name):
        try:
            X = self.data.drop(columns=self.target_column)
            y = self.data[self.target_column]

            if not self.numerical_data.empty:
                mode_values = self.data.mode().iloc[0]
                fig = px.bar(x=self.data.columns, y=mode_values)
                fig.update_layout(
                    title="Mode of Each Column",
                    xaxis_title="Columns",
                    yaxis_title="Mode Values",
                )
                fig.write_html(
                    os.path.join(
                        self.output_folder,
                        f"mode_plot_{os.path.basename(file_name).split('.')[0]}.html",
                    )
                )
                self.document.add_paragraph(f"Mode plot saved at {self.output_folder}")

            else:
                # Fallback if no numerical data
                self.document.add_paragraph("No numerical data to plot mode.")
        except Exception as e:
            logging.error(f"Error in plot_graphs: {e}")

    def save_documentation(self, file_name):
        try:
            doc_path = os.path.join(
                self.output_folder,
                "analysis_on_" + os.path.basename(file_name).split(".")[0],
            )
            os.makedirs(doc_path, exist_ok=True)
            self.document.save(os.path.join(doc_path, "analysis_report.docx"))
            print(
                f"Documentation saved at {os.path.join(doc_path, 'analysis_report.docx')}"
            )
        except Exception as e:
            logging.error(f"Error in save_documentation: {e}")

    def save_pickles(self, file_name):
        try:
            pkl_path = os.path.join(
                self.output_folder, "data_" + os.path.basename(file_name).split(".")[0]
            )
            os.makedirs(pkl_path, exist_ok=True)
            for model, _ in self.model_results:
                with open(
                    os.path.join(pkl_path, f"{model.__class__.__name__}.pkl"), "wb"
                ) as f:
                    pickle.dump(model, f)
            print(f"Pickles saved at {pkl_path}")
        except Exception as e:
            logging.error(f"Error in save_pickles: {e}")


# Usage
if __name__ == "__main__":
    try:
        analyser = Analyser()
        analyser.add_requirement()
        analyser.run_analysis()
    except Exception as e:
        logging.error(f"Critical error in main execution: {e}")
