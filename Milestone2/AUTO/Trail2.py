import os
import shutil
import re
import pickle
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import (
    StandardScaler,
    LabelEncoder,
    OneHotEncoder,
    MinMaxScaler,
    PolynomialFeatures,
)
from sklearn.linear_model import LinearRegression, LogisticRegression, Ridge, Lasso
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import (
    RandomForestClassifier,
    GradientBoostingClassifier,
    AdaBoostClassifier,
    VotingClassifier,
)
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import GaussianNB
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.metrics import accuracy_score, mean_squared_error, r2_score
from plotly import express as px
from plotly import graph_objects as go
from docx import Document
import logging
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Ensure the errors_log directory exists
log_dir = "root/errors_log"
os.makedirs(log_dir, exist_ok=True)

# Setting up logging
logging.basicConfig(
    filename=os.path.join(log_dir, "errors.log"),
    level=logging.ERROR,
    format="%(asctime)s:%(levelname)s:%(message)s",
)


class Analyser:
    def __init__(self):
        self.dataset_path = None
        self.files = []
        self.data = None
        self.document = Document()
        self.model_results = []
        self.best_model = None
        self.output_folder = None

    def add_requirement(self):
        try:
            self.dataset_path = input(
                "Enter the path of the dataset (CSV/XLSX file or folder containing them): "
            ).strip()
            self.output_folder = input("Enter the path of the output folder: ").strip()
            os.makedirs(self.output_folder, exist_ok=True)
            print(f"Dataset path set to: {self.dataset_path}")
        except Exception as e:
            logging.error(f"Error in add_requirement: {e}")

    def run_analysis(self):
        try:
            self.collect_files(self.dataset_path)
            if not self.files:
                print("No valid files found.")
            else:
                for file in self.files:
                    print(f"Analyzing file: {file}")
                    self.document.add_heading(f"Analysis for {file}", level=1)
                    self.clean_data(file)
                    self.select_target_column()
                    self.apply_models()
                    self.plot_graphs(file)
                    self.save_documentation(file)
                    self.save_pickles(file)
        except Exception as e:
            logging.error(f"Error in run_analysis: {e}")

    def collect_files(self, path):
        try:
            if os.path.isfile(path) and path.endswith((".csv", ".xlsx")):
                self.files.append(path)
            elif os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for file in files:
                        if file.endswith((".csv", ".xlsx")):
                            self.files.append(os.path.join(root, file))
            else:
                print(f"Invalid path: {path}")
        except Exception as e:
            logging.error(f"Error in collect_files: {e}")

    def clean_data(self, file_path):
        try:
            if file_path.endswith(".csv"):
                self.data = pd.read_csv(file_path)
            elif file_path.endswith(".xlsx"):
                self.data = pd.read_excel(file_path)

            # Split data into numerical and categorical types
            self.numerical_data = self.data.select_dtypes(include=["number"])
            self.categorical_data = self.data.select_dtypes(exclude=["number"])

            # Check for date and year types
            date_columns = [
                col
                for col in self.data.columns
                if "date" in col.lower()
                or pd.api.types.is_datetime64_any_dtype(self.data[col])
            ]
            year_columns = [col for col in self.data.columns if "year" in col.lower()]
            self.date_data = self.data[date_columns] if date_columns else pd.DataFrame()
            self.year_data = self.data[year_columns] if year_columns else pd.DataFrame()

            self.document.add_paragraph("Data Cleaning Steps:")
            self.document.add_paragraph(
                f"Numerical Data:\n{self.numerical_data.head().to_string()}"
            )
            self.document.add_paragraph(
                f"Categorical Data:\n{self.categorical_data.head().to_string()}"
            )
            self.document.add_paragraph(
                f"Date Data:\n{self.date_data.head().to_string()}"
            )
            self.document.add_paragraph(
                f"Year Data:\n{self.year_data.head().to_string()}"
            )
        except Exception as e:
            logging.error(f"Error in clean_data: {e}")

    def select_target_column(self):
        try:
            vectorizer = CountVectorizer().fit_transform(self.data.columns)
            vectors = vectorizer.toarray()
            csim = cosine_similarity(vectors)

            target_column = None
            best_match_score = 0

            for i, col1 in enumerate(self.data.columns):
                for j, col2 in enumerate(self.data.columns):
                    if i != j and "target" in col2.lower():
                        match_score = csim[i][j]
                        if match_score > best_match_score:
                            best_match_score = match_score
                            target_column = col2

            if target_column:
                print(f"Suggested target column based on NLP: {target_column}")
            else:
                target_column = self.user_select_target_column()

            user_confirm = (
                input(f"Is '{target_column}' the correct target column? (yes/no): ")
                .strip()
                .lower()
            )
            if user_confirm == "no":
                target_column = self.user_select_target_column()

            self.target_column = target_column
            print(f"Target column set to: {self.target_column}")
            self.document.add_paragraph(f"Selected target column: {self.target_column}")

        except Exception as e:
            logging.error(f"Error in select_target_column: {e}")

    def user_select_target_column(self):
        try:
            print("Available columns:")
            for idx, col in enumerate(self.data.columns):
                print(f"{idx}: {col}")

            user_selection = int(
                input(
                    "Select the target column by entering the corresponding number: "
                ).strip()
            )
            return self.data.columns[user_selection]
        except Exception as e:
            logging.error(f"Error in user_select_target_column: {e}")
            return None

    def transform_data(self, model, X, y=None):
        try:
            if isinstance(
                model,
                (
                    LinearRegression,
                    Ridge,
                    Lasso,
                    LogisticRegression,
                    DecisionTreeClassifier,
                    RandomForestClassifier,
                    GradientBoostingClassifier,
                    AdaBoostClassifier,
                    SVC,
                    KNeighborsClassifier,
                ),
            ):
                scaler = StandardScaler()
                X_transformed = scaler.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, KMeans):
                scaler = MinMaxScaler()
                X_transformed = scaler.fit_transform(X)
                return X_transformed
            elif isinstance(model, GaussianNB):
                label_encoder = LabelEncoder()
                for col in X.select_dtypes(include=["object"]).columns:
                    X[col] = label_encoder.fit_transform(X[col])
                return X, y
            elif isinstance(model, VotingClassifier):
                encoder = OneHotEncoder()
                X_transformed = encoder.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, PolynomialFeatures):
                poly = PolynomialFeatures(degree=2)
                X_transformed = poly.fit_transform(X)
                return X_transformed, y
            elif isinstance(model, PCA):
                pca = PCA(n_components=2)
                X_transformed = pca.fit_transform(X)
                return X_transformed, y
            return X, y
        except Exception as e:
            logging.error(f"Error in transform_data: {e}")

    def apply_models(self):
        try:
            X = self.data.drop(columns=self.target_column)
            y = self.data[self.target_column]

            models = [
                LinearRegression(),
                Ridge(),
                Lasso(),
                LogisticRegression(),
                DecisionTreeClassifier(),
                RandomForestClassifier(),
                GradientBoostingClassifier(),
                AdaBoostClassifier(),
                SVC(),
                KNeighborsClassifier(),
                GaussianNB(),
                KMeans(n_clusters=3),
                VotingClassifier(
                    estimators=[
                        ("lr", LogisticRegression()),
                        ("rf", RandomForestClassifier()),
                        ("svc", SVC()),
                    ]
                ),
                PolynomialFeatures(),
                PCA(n_components=2),
            ]

            best_score = float("-inf")

            for model in models:
                X_transformed, y_transformed = self.transform_data(model, X, y)
                model.fit(X_transformed, y_transformed)
                y_pred = model.predict(X_transformed)
                metrics = {}
                if isinstance(model, (LinearRegression, Ridge, Lasso)):
                    mse = mean_squared_error(y_transformed, y_pred)
                    r2 = r2_score(y_transformed, y_pred)
                    metrics = {"mse": mse, "r2": r2}
                    score = r2
                else:
                    accuracy = accuracy_score(y_transformed, y_pred)
                    metrics = {"accuracy": accuracy}
                    score = accuracy

                self.model_results.append((model, metrics))

                self.document.add_paragraph(
                    f"Model {model.__class__.__name__} metrics: {metrics}"
                )

                if score > best_score:
                    best_score = score
                    self.best_model = model

            self.document.add_paragraph(
                f"Best Model: {self.best_model.__class__.__name__} with score: {best_score}"
            )
            with open("best_model.pkl", "wb") as f:
                pickle.dump(self.best_model, f)
        except Exception as e:
            logging.error(f"Error in apply_models: {e}")

    def plot_graphs(self, file_name):
        try:
            X = self.data.drop(columns=self.target_column)
            y = self.data[self.target_column]

            # Check for numerical values
            if not self.numerical_data.empty:
                # Use mode method for numerical data
                mode_values = self.data.mode().iloc[0]
                fig = px.bar(x=self.data.columns, y=mode_values)
                fig.update_layout(
                    title="Mode of Each Column",
                    xaxis_title="Columns",
                    yaxis_title="Mode Values",
                )
                fig.write_html(
                    os.path.join(
                        self.output_folder,
                        f"mode_plot_{os.path.basename(file_name).split('.')[0]}.html",
                    )
                )
                self.document.add_paragraph(
                    f"Mode plot saved at: {os.path.join(self.output_folder, 'mode_plot_' + os.path.basename(file_name).split('.')[0] + '.html')}"
                )

            # Check if the dataset consists of only text values
            elif self.categorical_data.shape[1] > 0:
                print("The dataset consists of only text values.")
                x_col = input("Enter the name of the column to use for the x-axis: ")
                y_col = input(
                    "Enter the name of the column to use for the y-axis (or press Enter to use counts): "
                )

                if y_col.strip() == "":
                    # If no y_col provided, count occurrences
                    counts = self.data[x_col].value_counts()
                    fig = px.bar(counts, title=f"Counts of {x_col}")
                else:
                    # If y_col is provided
                    fig = px.bar(
                        self.data, x=x_col, y=y_col, title=f"{y_col} vs {x_col}"
                    )

                fig.write_html(
                    os.path.join(
                        self.output_folder,
                        f"text_plot_{os.path.basename(file_name).split('.')[0]}.html",
                    )
                )
                self.document.add_paragraph(
                    f"Text plot saved at: {os.path.join(self.output_folder, 'text_plot_' + os.path.basename(file_name).split('.')[0] + '.html')}"
                )
            else:
                print("No valid data to plot.")

            # Save plots as SVG for document
            svg_path = os.path.join(
                self.output_folder,
                f"plot_{os.path.basename(file_name).split('.')[0]}.svg",
            )
            fig.write_image(svg_path)
            self.document.add_picture(svg_path, width=go.layout.width(640))

            # Explanation of the plot
            self.document.add_paragraph(
                "This plot visualizes the data, showcasing either the mode of numerical columns or the distribution "
                "of text values."
            )

        except Exception as e:
            logging.error(f"Error in plot_graphs: {e}")

    def save_documentation(self, file_name):
        try:
            doc_path = os.path.join(
                self.output_folder,
                "analysis_on_" + os.path.basename(file_name).split(".")[0],
            )
            os.makedirs(doc_path, exist_ok=True)
            self.document.save(os.path.join(doc_path, "analysis_report.docx"))
            print(
                f"Documentation saved at {os.path.join(doc_path, 'analysis_report.docx')}"
            )
        except Exception as e:
            logging.error(f"Error in save_documentation: {e}")

    def save_pickles(self, file_name):
        try:
            pkl_path = os.path.join(
                self.output_folder, "data_" + os.path.basename(file_name).split(".")[0]
            )
            os.makedirs(pkl_path, exist_ok=True)
            for model, _ in self.model_results:
                with open(
                    os.path.join(pkl_path, f"{model.__class__.__name__}.pkl"), "wb"
                ) as f:
                    pickle.dump(model, f)
            print(f"Pickle files saved at {pkl_path}")
        except Exception as e:
            logging.error(f"Error in save_pickles: {e}")


if __name__ == "__main__":
    try:
        analyser = Analyser()
        analyser.add_requirement()
        analyser.run_analysis()
    except Exception as e:
        logging.error(f"Critical error in main execution: {e}")
